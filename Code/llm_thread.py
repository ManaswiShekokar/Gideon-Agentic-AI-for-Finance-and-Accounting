import contextlib
import json
import logging
import base64
import requests
from datetime import datetime, date
import io
import uuid
import calendar

import emoji
import markdown2

from odoo import _, api, fields, models
from odoo.exceptions import UserError
from psycopg2 import OperationalError

_logger = logging.getLogger(__name__)
import re
from collections import Counter



class LLMDocumentChunk(models.Model):
    """Model to store document chunks for RAG"""
    _name = 'llm.document.chunk'
    _description = 'LLM Document Chunk for RAG'
    _order = 'chunk_index'

    attachment_id = fields.Many2one('ir.attachment', string='Source Document', required=True, ondelete='cascade')
    chunk_text = fields.Text('Chunk Content', required=True)
    chunk_index = fields.Integer('Chunk Index', required=True)
    chunk_size = fields.Integer('Chunk Size')
    # Store ChromaDB document ID instead of raw embedding
    chroma_doc_id = fields.Char(string="ChromaDB ID", index=True)
    thread_id = fields.Many2one('llm.thread', string='Related Thread')
    processing_status = fields.Selection([
        ('pending', 'Pending'),
        ('processing', 'Processing'),
        ('completed', 'Completed'),
        ('failed', 'Failed')
    ], default='pending', string='Processing Status')
    
    @api.model
    def create_chunks_from_attachment(self, attachment_id, thread_id=None):
        """Create chunks from an attachment for RAG processing"""
        attachment = self.env['ir.attachment'].browse(attachment_id)
        if not attachment.exists():
            _logger.error(f"Attachment {attachment_id} not found")
            return False
            
        try:
            _logger.info(f"Starting chunk creation for {attachment.name}")
            
            # Extract text content
            text_content = self._extract_full_text(attachment)
            if not text_content:
                _logger.error(f"No text content extracted from {attachment.name}")
                return False
            
            _logger.info(f"Extracted {len(text_content)} characters from {attachment.name}")
            
            # Create chunks
            chunks = self._create_text_chunks(text_content, chunk_size=500, overlap=50)
            _logger.info(f"Created {len(chunks)} text chunks")
            
            # Delete existing chunks for this attachment PROPERLY
            existing_chunks = self.search([('attachment_id', '=', attachment_id), ('thread_id', '=', thread_id)])

            if existing_chunks:
                # Delete from ChromaDB first
                rag_service = self.env['llm.rag.service']
                collection_name = f"thread_{thread_id}" if thread_id else "general"
                for chunk in existing_chunks:
                    if chunk.chroma_doc_id:
                        rag_service._delete_from_chromadb(collection_name, chunk.chroma_doc_id)
                existing_chunks.unlink()
                _logger.info(f"Removed {len(existing_chunks)} existing chunks")
            
            # Create new chunks and add to ChromaDB
            chunk_records = []
            collection_name = f"thread_{thread_id}" if thread_id else "general"
            rag_service = self.env['llm.rag.service']
            
            for i, chunk_text in enumerate(chunks):
                # Generate unique document ID for ChromaDB
                chroma_doc_id = f"doc_{attachment_id}_chunk_{i}_{uuid.uuid4().hex[:8]}"
                
                chunk_vals = {
                    'attachment_id': attachment_id,
                    'chunk_text': chunk_text,
                    'chunk_index': i,
                    'chunk_size': len(chunk_text),
                    'thread_id': thread_id,
                    'chroma_doc_id': chroma_doc_id,
                    'processing_status': 'pending',
                }
                chunk_record = self.create(chunk_vals)
                chunk_records.append(chunk_record)
                
                # Add to ChromaDB with metadata
                metadata = {
                    'document_name': attachment.name,
                    'chunk_index': i,
                    'thread_id': thread_id or 0,
                    'attachment_id': attachment_id,
                    'file_type': attachment.name.split('.')[-1].lower() if '.' in attachment.name else 'unknown'
                }
                
                success = rag_service._add_to_chromadb(
                    collection_name=collection_name,
                    document_id=chroma_doc_id,
                    text=chunk_text,
                    metadata=metadata
                )
                
                if success:
                    chunk_record.processing_status = 'completed'
                    _logger.debug(f"Added chunk {i} to ChromaDB successfully")
                else:
                    chunk_record.processing_status = 'failed'
                    _logger.warning(f"Failed to add chunk {i} to ChromaDB")
            
            successful_chunks = [c for c in chunk_records if c.processing_status == 'completed']
            _logger.info(f"Successfully created {len(successful_chunks)}/{len(chunk_records)} chunks in ChromaDB")
            
            return chunk_records if chunk_records else False
            
        except Exception as e:
            _logger.error(f"Error creating chunks from attachment {attachment.name}: {str(e)}", exc_info=True)
            return False
    
    def _extract_full_text(self, attachment):
        """Extract full text from attachment based on file type"""
        try:
            if not attachment.datas:
                _logger.error(f"No data in attachment {attachment.name}")
                return None
                
            file_data = base64.b64decode(attachment.datas)
            file_ext = attachment.name.split('.')[-1].lower() if '.' in attachment.name else ''
            
            _logger.info(f"Processing file {attachment.name} with extension {file_ext}")
            
            if file_ext in ['txt', 'md']:
                try:
                    return file_data.decode('utf-8')
                except UnicodeDecodeError:
                    return file_data.decode('latin-1')
            
            elif file_ext == 'csv':
                try:
                    return file_data.decode('utf-8')
                except UnicodeDecodeError:
                    return file_data.decode('latin-1')
            
            elif file_ext == 'pdf':
                return self._extract_pdf_text(file_data)
            
            elif file_ext in ['doc', 'docx']:
                return self._extract_docx_text(file_data)
            
            elif file_ext in ['xlsx', 'xls']:
                return self._extract_excel_text(file_data)
            
            else:
                _logger.warning(f"Unsupported file type: {file_ext}")
                # Try to decode as text anyway
                try:
                    return file_data.decode('utf-8')
                except:
                    return None
                
        except Exception as e:
            _logger.error(f"Error extracting text from {attachment.name}: {str(e)}", exc_info=True)
            return None
    
    def _extract_pdf_text(self, file_data):
        """Extract text from PDF with fallback options"""
        # Try PyPDF2 first
        try:
            import PyPDF2
            pdf_file = io.BytesIO(file_data)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            if text.strip():
                return text.strip()
        except ImportError:
            _logger.warning("PyPDF2 not installed")
        except Exception as e:
            _logger.error(f"PyPDF2 extraction failed: {str(e)}")
        
        # Try pdfplumber as fallback
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_data)) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                
                if text.strip():
                    return text.strip()
        except ImportError:
            _logger.warning("pdfplumber not installed")
        except Exception as e:
            _logger.error(f"pdfplumber extraction failed: {str(e)}")
        
        _logger.error("All PDF extraction methods failed")
        return None
    
    def _extract_docx_text(self, file_data):
        """Extract text from DOCX using python-docx"""
        try:
            from docx import Document
            doc_file = io.BytesIO(file_data)
            doc = Document(doc_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip() if text.strip() else None
        except ImportError:
            _logger.warning("python-docx not installed. Cannot extract DOCX text.")
            return None
        except Exception as e:
            _logger.error(f"Error extracting DOCX text: {str(e)}", exc_info=True)
            return None
    
    def _extract_excel_text(self, file_data):
        """Extract text from Excel using pandas or openpyxl"""
        # Try pandas first
        try:
            import pandas as pd
            excel_file = io.BytesIO(file_data)
            
            # Read all sheets
            all_sheets = pd.read_excel(excel_file, sheet_name=None)
            
            text = ""
            for sheet_name, df in all_sheets.items():
                text += f"Sheet: {sheet_name}\n"
                text += df.to_string(index=False) + "\n\n"
            
            return text.strip() if text.strip() else None
        except ImportError:
            _logger.warning("pandas not installed")
        except Exception as e:
            _logger.error(f"pandas extraction failed: {str(e)}")
        
        # Try openpyxl as fallback
        try:
            import openpyxl
            excel_file = io.BytesIO(file_data)
            workbook = openpyxl.load_workbook(excel_file, data_only=True)
            
            text = ""
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text += f"Sheet: {sheet_name}\n"
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = "\t".join([str(cell) if cell is not None else "" for cell in row])
                    text += row_text + "\n"
                text += "\n"
            
            return text.strip() if text.strip() else None
        except ImportError:
            _logger.warning("openpyxl not installed")
        except Exception as e:
            _logger.error(f"openpyxl extraction failed: {str(e)}")
        
        return None
    
    def _create_text_chunks(self, text, chunk_size=500, overlap=50):
        """Split text into overlapping chunks with better sentence handling"""
        if not text:
            return []
            
        # First try to split by sentences
        import re
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        if len(sentences) <= 1:
            # Fallback to word-based chunking
            return self._create_word_chunks(text, chunk_size, overlap)
        
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            # If adding this sentence would exceed chunk size
            if len(current_chunk) + len(sentence) > chunk_size and current_chunk:
                chunks.append(current_chunk.strip())
                # Start new chunk with overlap from previous chunk
                overlap_text = current_chunk[-overlap:] if len(current_chunk) > overlap else current_chunk
                current_chunk = overlap_text + " " + sentence
            else:
                current_chunk += " " + sentence if current_chunk else sentence
        
        # Add the last chunk
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def _create_word_chunks(self, text, chunk_size=500, overlap=50):
        """Fallback word-based chunking"""
        words = text.split()
        chunks = []
        
        i = 0
        while i < len(words):
            chunk_words = words[i:i + chunk_size]
            chunk_text = ' '.join(chunk_words)
            chunks.append(chunk_text)
            
            i += max(1, chunk_size - overlap)
            
            if i >= len(words):
                break
        
        return chunks


class LLMRAGService(models.Model):
    """RAG Service for document retrieval and LLM integration using ChromaDB"""
    _name = 'llm.rag.service'
    _description = 'RAG Service for LLM with ChromaDB'
    
    def _get_chromadb_client(self):
        """Get ChromaDB client with v2 API compatibility"""
        try:
            import chromadb
            from chromadb.config import Settings
            
            # Connect to ChromaDB server running in Docker with v2 API settings
            client = chromadb.HttpClient(
                host="chroma", 
                port=8000,
                settings=Settings(
                    anonymized_telemetry=False,
                    allow_reset=True
                )
            )
            return client
        except ImportError:
            _logger.error("ChromaDB client not installed. Please install: pip install chromadb")
            return None
        except Exception as e:
            _logger.error(f"Failed to connect to ChromaDB: {str(e)}")
            return None
    
    def _get_or_create_collection(self, collection_name):
        """Get or create a ChromaDB collection with v2 API"""
        try:
            client = self._get_chromadb_client()
            if not client:
                return None
            
            # Try to get existing collection first
            try:
                collection = client.get_collection(name=collection_name)
                _logger.debug(f"Retrieved existing collection: {collection_name}")
                return collection
            except Exception:
                # Collection doesn't exist, create it with v2 compatible settings
                try:
                    collection = client.create_collection(
                        name=collection_name,
                        metadata={"description": f"RAG collection for {collection_name}"},
                        # Use default embedding function for v2
                        embedding_function=None  # Will use default
                    )
                    _logger.info(f"Created new collection: {collection_name}")
                    return collection
                except Exception as e:
                    _logger.error(f"Failed to create collection {collection_name}: {str(e)}")
                    return None
                
        except Exception as e:
            _logger.error(f"Error getting/creating collection {collection_name}: {str(e)}")
            return None
    
    def _add_to_chromadb(self, collection_name, document_id, text, metadata):
        """Add document to ChromaDB with v2 API"""
        try:
            collection = self._get_or_create_collection(collection_name)
            if not collection:
                return False
            
            # For v2 API, we can either provide embeddings or let ChromaDB generate them
            # Let's use Ollama for consistency with your existing setup
            embedding = self._generate_embedding_with_ollama(text)
            if not embedding:
                _logger.error(f"Failed to generate embedding for document {document_id}")
                return False
            
            # Add to ChromaDB - v2 API method
            collection.add(
                documents=[text],
                embeddings=[embedding],
                metadatas=[metadata],
                ids=[document_id]
            )
            
            _logger.debug(f"Added document {document_id} to ChromaDB collection {collection_name}")
            return True
            
        except Exception as e:
            _logger.error(f"Error adding document to ChromaDB: {str(e)}")
            return False
    
    def _delete_from_chromadb(self, collection_name, document_id):
        """Delete document from ChromaDB with v2 API"""
        try:
            collection = self._get_or_create_collection(collection_name)
            if not collection:
                return False
            
            # v2 API delete method
            collection.delete(ids=[document_id])
            _logger.debug(f"Deleted document {document_id} from ChromaDB")
            return True
            
        except Exception as e:
            _logger.error(f"Error deleting document from ChromaDB: {str(e)}")
            return False
    
    def _preprocess_text_for_embedding(self, text):
        """Preprocess text to improve embedding quality"""
        try:
            # Remove excessive whitespace
            text = ' '.join(text.split())
            
            # Truncate if too long (most embedding models have token limits)
            max_chars = 2000  # Adjust based on your embedding model
            if len(text) > max_chars:
                text = text[:max_chars] + "..."
            
            return text
        except Exception as e:
            _logger.error(f"Error preprocessing text: {str(e)}")
            return text
        
    def _normalize_embedding(self, embedding):
        """Normalize embedding vector to unit length"""
        try:
            import math
            
            # Calculate the magnitude (L2 norm)
            magnitude = math.sqrt(sum(x * x for x in embedding))
            
            if magnitude > 0:
                # Normalize to unit vector
                return [x / magnitude for x in embedding]
            else:
                return embedding
                
        except Exception as e:
            _logger.error(f"Error normalizing embedding: {str(e)}")
            return embedding

    def _generate_embedding_with_ollama(self, text, max_retries=2):
        """Generate embedding using Ollama with improved error handling and model fallback"""
        try:
            # Preprocess text for better embeddings
            processed_text = self._preprocess_text_for_embedding(text)
            
            # Try multiple embedding models with different strategies
            embedding_strategies = [
                {'model': 'nomic-embed-text:latest', 'normalize': True},
                {'model': 'all-minilm:latest', 'normalize': True},
                {'model': 'mxbai-embed-large:latest', 'normalize': False},
            ]
            
            for strategy in embedding_strategies:
                for retry in range(max_retries):
                    try:
                        response = requests.post(
                            'http://ollama:11434/api/embed',
                            json={
                                'model': strategy['model'],
                                'input': processed_text
                            },
                            timeout=60
                        )
                        
                        if response.status_code == 200:
                            embedding_data = response.json()
                            embeddings = embedding_data.get('embeddings', [])
                            
                            if embeddings and len(embeddings) > 0:
                                embedding = embeddings[0]
                                
                                # Normalize embedding if specified
                                if strategy['normalize']:
                                    embedding = self._normalize_embedding(embedding)
                                
                                _logger.info(f"Generated embedding using {strategy['model']}, length: {len(embedding)}")
                                return embedding
                        
                    except requests.exceptions.RequestException as e:
                        _logger.warning(f"Ollama request failed for model {strategy['model']} (attempt {retry + 1}): {str(e)}")
                        if retry < max_retries - 1:
                            import time
                            time.sleep(1)  # Wait before retry
            
            _logger.error("All embedding generation attempts failed")
            return None 
            
        except Exception as e:
            _logger.error(f"Error generating embedding with Ollama: {str(e)}")
            return None
    
    def search_relevant_chunks(self, query, thread_id=None, limit=5, min_similarity=0.1):
        """Search for relevant document chunks using ChromaDB v2 API with improved similarity calculation"""
        try:
            collection_name = f"thread_{thread_id}" if thread_id else "general"
            collection = self._get_or_create_collection(collection_name)
            
            if not collection:
                _logger.warning(f"Collection {collection_name} not found")
                return []
            
            # Generate embedding for the query
            query_embedding = self._generate_embedding_with_ollama(query)
            if not query_embedding:
                _logger.warning("Failed to generate query embedding")
                return []
            
            # Search in ChromaDB with v2 API - increase n_results for better matching
            results = collection.query(
                query_embeddings=[query_embedding],
                n_results=min(limit * 3, 20),  # Get more results to filter better ones
                include=['documents', 'metadatas', 'distances']
            )
            
            # Process results with improved similarity calculation
            relevant_chunks = []
            if results and 'documents' in results and results['documents']:
                documents = results['documents'][0] if results['documents'] else []
                metadatas = results['metadatas'][0] if results.get('metadatas') else []
                distances = results['distances'][0] if results.get('distances') else []
                
                for i, doc in enumerate(documents):
                    metadata = metadatas[i] if i < len(metadatas) else {}
                    distance = distances[i] if i < len(distances) else 1.0
                    
                    # Improved similarity calculation
                    # ChromaDB typically returns cosine distances, convert to similarity
                    if distance <= 0:
                        similarity = 1.0  # Perfect match
                    else:
                        # For cosine distance: similarity = 1 - distance
                        # Clamp between 0 and 1
                        similarity = max(0, min(1, 1 - distance))
                    
                    # Apply similarity boost based on query matching
                    similarity = self._boost_similarity(doc, query, similarity)
                    
                    # Only include chunks above minimum similarity threshold
                    if similarity >= min_similarity:
                        relevant_chunks.append({
                            'content': doc,
                            'similarity': similarity,
                            'distance': distance,  # Keep original distance for debugging
                            'document_name': metadata.get('document_name', 'Unknown'),
                            'chunk_index': metadata.get('chunk_index', 0),
                            'attachment_id': metadata.get('attachment_id', 0)
                        })
            
            # Sort by similarity (highest first) and limit results
            relevant_chunks = sorted(relevant_chunks, key=lambda x: x['similarity'], reverse=True)[:limit]
            
            _logger.info(f"Found {len(relevant_chunks)} relevant chunks for query")
            for i, chunk in enumerate(relevant_chunks):
                _logger.info(f"  Chunk {i}: similarity={chunk['similarity']:.3f}, distance={chunk.get('distance', 0):.3f}")
            
            return relevant_chunks
            
        except Exception as e:
            _logger.error(f"Error searching relevant chunks: {str(e)}")
            return []
    
    def _boost_similarity(self, document_text, query, base_similarity):
        """Boost similarity based on keyword matching and other factors"""
        try:
            # Convert to lowercase for comparison
            doc_lower = document_text.lower()
            query_lower = query.lower()
            
            # Keyword matching boost
            query_words = set(query_lower.split())
            doc_words = set(doc_lower.split())
            
            # Calculate word overlap
            common_words = query_words.intersection(doc_words)
            if query_words:
                word_overlap_ratio = len(common_words) / len(query_words)
                
                # Boost similarity based on word overlap
                if word_overlap_ratio > 0.5:  # More than 50% words match
                    base_similarity = min(1.0, base_similarity + 0.2)
                elif word_overlap_ratio > 0.3:  # More than 30% words match
                    base_similarity = min(1.0, base_similarity + 0.1)
                elif word_overlap_ratio > 0.1:  # More than 10% words match
                    base_similarity = min(1.0, base_similarity + 0.05)
            
            # Exact phrase matching boost
            if query_lower in doc_lower:
                base_similarity = min(1.0, base_similarity + 0.15)
            
            # Length penalty for very short documents
            if len(document_text) < 50:
                base_similarity *= 0.8
            
            return base_similarity
            
        except Exception as e:
            _logger.error(f"Error boosting similarity: {str(e)}")
            return base_similarity

    def generate_rag_response(self, query, thread_id=None, context_limit=3, min_similarity=0.1):
        """Generate response using RAG + LLM with enhanced similarity handling"""
        try:
            _logger.info("="*30)
            _logger.info("RAG SERVICE DEBUG START")
            _logger.info(f"Query: '{query}'")
            _logger.info(f"Thread ID: {thread_id}")
            _logger.info(f"Context limit: {context_limit}")
            _logger.info(f"Min similarity threshold: {min_similarity}")
            
            # 1. Search for relevant document chunks with minimum similarity
            _logger.info("Step 1: Searching for relevant chunks...")
            relevant_chunks = self.search_relevant_chunks(
                query, thread_id, 
                limit=context_limit, 
                min_similarity=min_similarity
            )
            _logger.info(f"Found {len(relevant_chunks)} relevant chunks")
            
            for i, chunk in enumerate(relevant_chunks):
                similarity_pct = chunk.get('similarity', 0) * 100
                _logger.info(f"  Chunk {i}: doc='{chunk.get('document_name', 'Unknown')}', similarity={similarity_pct:.1f}%")
            
            if not relevant_chunks:
                _logger.warning("No relevant chunks found above similarity threshold")
                return {
                    'response': f"I don't have any documents with sufficient relevance (above {min_similarity*100:.1f}% similarity) to answer your question.",
                    'sources': [],
                    'min_similarity_used': min_similarity
                }
            
            # 2. Build enhanced context from relevant chunks
            _logger.info("Step 2: Building enhanced context from chunks...")
            context_parts = []
            sources = []
            
            for i, chunk_data in enumerate(relevant_chunks):
                doc_name = chunk_data.get('document_name', f'Document {i+1}')
                content = chunk_data.get('content', '')
                similarity = chunk_data.get('similarity', 0)
                
                # Add similarity score to context for better LLM awareness
                context_parts.append(f"Document {i+1} ({doc_name}) [Relevance: {similarity*100:.1f}%]:\n{content}")
                sources.append({
                    'document': doc_name,
                    'similarity': similarity,
                    'chunk_index': chunk_data.get('chunk_index', 0),
                    'attachment_id': chunk_data.get('attachment_id', 0)
                })
            
            context = "\n\n".join(context_parts)
            _logger.info(f"Enhanced context built: {len(context)} characters")
            
            # 3. Create enhanced prompt with similarity awareness
            enhanced_prompt = f"""Based on the following document excerpts (with relevance scores), please answer the user's question. Higher relevance scores indicate better matches to the query.

    Document Context:
    {context}

    User Question: {query}

    Instructions:
    - Use the most relevant documents (higher relevance scores) as primary sources
    - If multiple documents mention the same information, synthesize the response
    - If the information seems contradictory between documents, note this
    - Provide a comprehensive answer based on the available information
    - If you reference specific information, mention which document it came from"""

            _logger.info(f"Enhanced prompt created: {len(enhanced_prompt)} characters")
            
            # 4. Send to Ollama for response with longer timeout for complex queries
            _logger.info("Step 3: Sending to Ollama...")
            response = requests.post(
                'http://ollama:11434/api/generate',
                json={
                    'model': 'llama3.1:8b',
                    'prompt': enhanced_prompt,
                    'stream': False,
                    'options': {
                        'temperature': 0.7,  # Slightly more creative
                        'top_p': 0.9,
                        'num_ctx': 4096,  # Larger context window
                    }
                },
                timeout=180  # Longer timeout for complex responses
            )
            
            _logger.info(f"Ollama response status: {response.status_code}")
            
            if response.status_code == 200:
                response_data = response.json()
                llm_response = response_data.get('response', 'No response generated')
                _logger.info(f"LLM response received: {len(llm_response)} characters")
                
                result = {
                    'response': llm_response,
                    'sources': sources,
                    'context_used': len(relevant_chunks),
                    'min_similarity_used': min_similarity,
                    'avg_similarity': sum(s['similarity'] for s in sources) / len(sources) if sources else 0
                }
                _logger.info("RAG SERVICE DEBUG END - SUCCESS")
                _logger.info("="*30)
                return result
            else:
                _logger.error(f"LLM generation failed: {response.status_code}")
                _logger.error(f"Response text: {response.text}")
                return {
                    'response': "Sorry, I encountered an error generating the response.",
                    'sources': sources,
                    'min_similarity_used': min_similarity
                }
                
        except Exception as e:
            _logger.error(f"Error generating RAG response: {str(e)}", exc_info=True)
            return {
                'response': f"Error generating response: {str(e)}",
                'sources': [],
                'min_similarity_used': min_similarity
            }
    
    def test_chromadb_connection(self):
        """Test ChromaDB connection using v2 API"""
        try:
            client = self._get_chromadb_client()
            if not client:
                return False, "Failed to create ChromaDB client"
            
            # Test with v2 heartbeat endpoint
            response = requests.get('http://chroma:8000/api/v2/heartbeat', timeout=10)
            if response.status_code == 200:
                return True, "ChromaDB connection successful"
            else:
                return False, f"ChromaDB heartbeat failed: {response.status_code}"
                
        except Exception as e:
            return False, f"ChromaDB connection test failed: {str(e)}"


class RelatedRecordProxy:
    """
    A proxy object that provides clean access to related record fields in Jinja templates.
    Usage in templates: {{ related_record.get_field('field_name', 'default_value') }}
    When called directly, returns JSON with model name, id, and display name.
    """

    def __init__(self, record):
        self._record = record

    def get_field(self, field_name, default=""):
        """
        Get a field value from the related record.

        Args:
            field_name (str): The field name to access
            default: Default value if field doesn't exist or is empty

        Returns:
            The field value, or default if not available
        """
        if not self._record:
            return default

        try:
            if hasattr(self._record, field_name):
                value = getattr(self._record, field_name)

                # Handle different field types
                if value is None:
                    return default
                elif isinstance(value, bool):
                    return value  # Keep as boolean for Jinja
                elif hasattr(value, "name"):  # Many2one field
                    return value.name
                elif hasattr(value, "mapped"):  # Many2many/One2many field
                    return value.mapped("name")
                else:
                    return value
            else:
                _logger.debug(
                    "Field '%s' not found on record %s", field_name, self._record
                )
                return default

        except Exception as e:
            _logger.error(
                "Error getting field '%s' from record: %s", field_name, str(e)
            )
            return default

    def __getattr__(self, name):
        """Allow direct attribute access as fallback"""
        return self.get_field(name)

    def __bool__(self):
        """Return True if we have a record"""
        return bool(self._record)

    def __str__(self):
        """When called by itself, return JSON of model name, id, and display name"""
        if not self._record:
            return json.dumps({"model": None, "id": None, "display_name": None})

        return json.dumps(
            {
                "model": self._record._name,
                "id": self._record.id,
                "display_name": getattr(
                    self._record, "display_name", str(self._record)
                ),
            }
        )

    def __repr__(self):
        """Same as __str__ for consistency"""
        return self.__str__()


class LLMThread(models.Model):
    _name = "llm.thread"
    _description = "LLM Chat Thread"
    _inherit = ["mail.thread"]
    _order = "write_date DESC"

    # ============================================================================
    # TABLE CONFIGURATION - Supported Tables Configuration  
    # ============================================================================

    SUPPORTED_TABLES = {
        'sale_order': {
            'model': 'sale.order',
            'name': 'Sales Orders',
            'key_fields': ['partner_id', 'date_order'],
            'required_fields': ['partner_id'],
            'default_values': {'state': 'draft'}
        },
        'purchase_order': {
            'model': 'purchase.order',
            'name': 'Purchase Orders',
            'key_fields': ['partner_id', 'date_order'],
            'required_fields': ['partner_id'],
            'default_values': {'state': 'draft'}
        },
        'stock_move': {
            'model': 'stock.move',
            'name': 'Stock Moves',
            'key_fields': ['product_id', 'location_id', 'location_dest_id'],
            'required_fields': ['product_id', 'product_uom_qty', 'location_id', 'location_dest_id'],
            'default_values': {'state': 'draft', 'name': 'Imported Stock Move'}
        },
        'stock_quant': {
            'model': 'stock.quant',
            'name': 'Stock Quantities',
            'key_fields': ['product_id', 'location_id'],
            'required_fields': ['product_id', 'location_id', 'quantity'],
            'default_values': {'reserved_quantity': 0}
        },
        'stock_valuation_layer': {
            'model': 'stock.valuation.layer',
            'name': 'Stock Valuation Layers',
            'key_fields': ['product_id'],
            'required_fields': ['product_id', 'quantity', 'value'],
            'default_values': {}
        },
        'hr_expense': {
            'model': 'hr.expense',
            'name': 'HR Expenses',
            'key_fields': ['employee_id', 'date'],
            'required_fields': ['employee_id', 'product_id', 'total_amount'],
            'default_values': {'state': 'draft'}
        },
        'hr_expense_sheet': {
            'model': 'hr.expense.sheet',
            'name': 'HR Expense Sheets',
            'key_fields': ['employee_id', 'accounting_date'],
            'required_fields': ['employee_id'],
            'default_values': {'state': 'draft'}
        },
        'account_move': {
            'model': 'account.move',
            'name': 'Journal Entries',
            'key_fields': ['date', 'journal_id'],
            'required_fields': ['journal_id'],
            'default_values': {'state': 'draft', 'move_type': 'entry'}
        }
    }

    name = fields.Char(
        string="Title",
        required=True,
    )
    user_id = fields.Many2one(
        "res.users",
        string="User",
        default=lambda self: self.env.user,
        required=True,
        ondelete="restrict",
    )
    provider_id = fields.Many2one(
        "llm.provider",
        string="Provider",
        required=True,
        ondelete="restrict",
    )
    model_id = fields.Many2one(
        "llm.model",
        string="Model",
        required=True,
        domain="[('provider_id', '=', provider_id), ('model_use', 'in', ['chat', 'multimodal'])]",
        ondelete="restrict",
    )
    active = fields.Boolean(default=True)

    # Updated fields for related record reference
    model = fields.Char(
        string="Related Document Model", help="Technical name of the related model"
    )
    res_id = fields.Many2oneReference(
        string="Related Document ID",
        model_field="model",
        help="ID of the related record",
    )

    tool_ids = fields.Many2many(
        "llm.tool",
        string="Available Tools",
        help="Tools that can be used by the LLM in this thread",
    )
    
    attachment_ids = fields.Many2many(
        'ir.attachment',
        string='All Thread Attachments',
        compute='_compute_attachment_ids',
        store=True,
        help='All attachments from all messages in this thread'
    )
    
    attachment_count = fields.Integer(
        string='Attachment Count',
        compute='_compute_attachment_count',
        store=True,
        help='Total number of attachments in this thread'
    )

    enable_expense_analysis = fields.Boolean(
        string="Enable Expense Analysis",
        default=False,
        help="Enable AI-powered expense analysis features"
    )

    pending_context = fields.Json(
        string="Pending Action Context",
        help="Stores context for multi-step operations like data appending confirmation."
    )
    
    # RAG-related fields
    document_chunk_ids = fields.One2many('llm.document.chunk', 'thread_id', string='Document Chunks')
    rag_enabled = fields.Boolean(string='RAG Enabled', default=True, help='Enable Retrieval-Augmented Generation for this thread')

    @api.model_create_multi
    def create(self, vals_list):
        """Set default title if not provided"""
        for vals in vals_list:
            if not vals.get("name"):
                vals["name"] = f"Chat with {self.model_id.name}"
        return super().create(vals_list)

    @api.depends('message_ids.attachment_ids')
    def _compute_attachment_ids(self):
        """Compute all attachments from all messages in this thread."""
        for thread in self:
            # Get all attachments from all messages in this thread
            all_attachments = thread.message_ids.mapped('attachment_ids')
            thread.attachment_ids = [(6, 0, all_attachments.ids)]
    
    @api.depends('attachment_ids')
    def _compute_attachment_count(self):
        """Compute the total number of attachments in this thread."""
        for thread in self:
            thread.attachment_count = len(thread.attachment_ids)

    def process_uploaded_documents(self):
        """Process all uploaded documents for RAG"""
        try:
            # Get all attachments for this thread
            attachments = self.env['ir.attachment'].search([
                ('res_model', '=', 'llm.thread'),
                ('res_id', '=', self.id)
            ])
            
            processed_count = 0
            for attachment in attachments:
                # Check if chunks already exist
                existing_chunks = self.env['llm.document.chunk'].search([
                    ('attachment_id', '=', attachment.id)
                ])
                
                if not existing_chunks:
                    chunks = self.env['llm.document.chunk'].create_chunks_from_attachment(
                        attachment.id, self.id
                    )
                    if chunks:
                        processed_count += 1
            
            _logger.info(f"Processed {processed_count} documents for thread {self.id}")
            return processed_count
            
        except Exception as e:
            _logger.error(f"Error processing documents for thread {self.id}: {str(e)}")
            return 0

    # ============================================================================
    # MESSAGE POST OVERRIDES - Clean integration with mail.thread
    # ============================================================================

    @api.returns("mail.message", lambda value: value.id)
    def message_post(self, *, llm_role=None, message_type="comment", **kwargs):
        """Override to handle LLM-specific message types and metadata.

        Args:
            llm_role (str): The LLM role ('user', 'assistant', 'tool', 'system')
                           If provided, will automatically set the appropriate subtype
        """

        # Convert LLM role to subtype_xmlid if provided
        if llm_role:
            _, role_to_id = self.env["mail.message"].get_llm_roles()
            if llm_role in role_to_id:
                # Get the xmlid from the role
                subtype_xmlid = f"llm.mt_{llm_role}"
                kwargs["subtype_xmlid"] = subtype_xmlid

        # Handle LLM-specific subtypes and email_from generation
        if not kwargs.get("author_id") and not kwargs.get("email_from"):
            kwargs["email_from"] = self._get_llm_email_from(
                kwargs.get("subtype_xmlid"), kwargs.get("author_id"), llm_role
            )

        # Convert markdown to HTML if needed (except for tool messages which use body_json)
        if kwargs.get("body") and llm_role != "tool":
            kwargs["body"] = self._process_llm_body(kwargs["body"])

        # Create the message using standard mail.thread flow
        message = super().message_post(message_type=message_type, **kwargs)
        
        # Auto-process attachments for RAG if enabled
        if self.rag_enabled and message.attachment_ids:
            self.process_uploaded_documents()
        
        return message

    def _get_llm_email_from(self, subtype_xmlid, author_id, llm_role=None):
        """Generate appropriate email_from for LLM messages."""
        if author_id:
            return None  # Let standard flow handle it

        provider_name = self.provider_id.name
        model_name = self.model_id.name

        if subtype_xmlid == "llm.mt_tool" or llm_role == "tool":
            return f"Tool <tool@{provider_name.lower().replace(' ', '')}.ai>"
        elif subtype_xmlid == "llm.mt_assistant" or llm_role == "assistant":
            return f"{model_name} <ai@{provider_name.lower().replace(' ', '')}.ai>"

        return None

    def _process_llm_body(self, body):
        """Process body content for LLM messages (markdown to HTML conversion)."""
        if not body:
            return body
        return markdown2.markdown(emoji.demojize(body))

    # ============================================================================
    # STREAMING MESSAGE CREATION
    # ============================================================================

    def message_post_from_stream(
        self, stream, llm_role, placeholder_text="…", **kwargs
    ):
        """Create and update a message from a streaming response.

        Args:
            stream: Generator yielding chunks of response data
            llm_role (str): The LLM role ('user', 'assistant', 'tool', 'system')
            placeholder_text (str): Text to show while streaming

        Returns:
            message: The created/updated message record
        """
        message = None
        accumulated_content = ""

        for chunk in stream:
            # Initialize message on first content
            if message is None and chunk.get("content"):
                message = self.message_post(
                    body=placeholder_text, llm_role=llm_role, author_id=False, **kwargs
                )
                yield {"type": "message_create", "message": message.message_format()[0]}

            # Handle content streaming
            if chunk.get("content") and message is not None:
                accumulated_content += chunk["content"]
                message.write({"body": self._process_llm_body(accumulated_content)})
                yield {"type": "message_chunk", "message": message.message_format()[0]}

            # Handle errors
            if chunk.get("error"):
                yield {"type": "error", "error": chunk["error"]}
                return message

        # Final update for assistant message
        if message and accumulated_content:
            message.write({"body": self._process_llm_body(accumulated_content)})
            yield {"type": "message_update", "message": message.message_format()[0]}

        return message
    
    # ============================================================================
    # GENERATION FLOW - Refactored to use message_post with roles
    # ============================================================================
    
    def get_latest_llm_message(self):
        """Get the most recent LLM message for flow control."""
        self.ensure_one()
        
        messages = self.env['mail.message'].search([
            ('model', '=', 'llm.thread'),
            ('res_id', '=', self.id),
            ('llm_role', 'in', ['user', 'assistant', 'tool'])
        ], order='create_date desc, id desc', limit=1)
        
        if not messages:
            return None
        return messages[0]

    def generate(self, user_message_body, **kwargs):
        """Main generation method with direct table handling and RAG fallback"""
        self.ensure_one()
        import re 
        with self._generation_lock():
            last_message = False
            # Post user message if provided
            if user_message_body:
                last_message = self.message_post(
                    body=user_message_body,
                    llm_role="user",
                    author_id=self.env.user.partner_id.id,
                    **kwargs,
                )
                yield {
                    "type": "message_create",
                    "message": last_message.message_format()[0],
                }

            # EXPENSE ANALYSIS - Handle expense-related queries first
            if self.enable_expense_analysis and self._is_expense_query(user_message_body):
                _logger.info("PROCESSING EXPENSE QUERY")
                try:
                    resp = self.generate_expense_response(user_message_body)
                    assistant_message = self.message_post(
                        body=resp,
                        llm_role="assistant",
                        author_id=False,
                        **kwargs,
                    )
                    _logger.info("EXPENSE ANALYSIS COMPLETED")
                    yield {"type": "message_create", "message": assistant_message.message_format()[0]}
                    return assistant_message  # IMPORTANT: prevents other paths from overriding
                except Exception as e:
                    _logger.error(f"Error in expense analysis: {e}")
    
                # PROCESS USER MESSAGE FOR SPECIAL HANDLING (Direct table requests, etc.)
                try:
                    # Extract clean text from HTML body for processing
                    clean_text = re.sub(r'<[^>]+>', '', user_message_body).strip()
                    
                    # ✅ Handle both boolean and message object returns
                    special_result = self._process_user_message(clean_text)
                    if special_result:
                        # If it's a message object, yield it
                        if hasattr(special_result, 'message_format'):
                            _logger.info("Special processing returned message - yielding it")
                            yield {"type": "message_create", "message": special_result.message_format()[0]}
                            return special_result
                        # If it's just True, special processing completed
                        else:
                            _logger.info("Special processing completed")
                            return last_message
                except Exception as e:
                    _logger.error(f"Error in _process_user_message: {e}")
                    # Continue with normal processing if there's an error

                # TRY RAG for semantic queries (non-table requests)
                if self.rag_enabled and last_message:
                    _logger.info("="*50)
                    _logger.info("RAG PROCESSING FOR SEMANTIC QUERIES")
                    
                    # Extract clean text from HTML body for RAG processing
                    clean_text = re.sub(r'<[^>]+>', '', last_message.body).strip()
                    _logger.info(f"Clean text for RAG: '{clean_text[:100]}...'")
                    
                    # Get completed chunks count
                    completed_chunks = self.document_chunk_ids.filtered(lambda c: c.processing_status == 'completed')
                    chunk_count = len(completed_chunks)
                    _logger.info(f"Available RAG chunks: {chunk_count}")
                    
                    if chunk_count > 0:
                        _logger.info("ATTEMPTING RAG RESPONSE")
                        
                        try:
                            # Use RAG service for semantic search
                            rag_service = self.env['llm.rag.service']
                            rag_result = rag_service.generate_rag_response(
                                clean_text, 
                                self.id, 
                                context_limit=5,
                                min_similarity=0.05  # Lower threshold for better matching
                            )
                            
                            if rag_result and isinstance(rag_result, dict):
                                response_text = rag_result.get('response', '').strip()
                                sources = rag_result.get('sources', [])
                                avg_similarity = rag_result.get('avg_similarity', 0)
                                
                                _logger.info(f"RAG response length: {len(response_text)}")
                                _logger.info(f"Average similarity: {avg_similarity:.3f}")
                                
                                # Check if we got a meaningful response
                                if response_text and len(response_text) > 20:
                                    _logger.info("RAG RESPONSE ACCEPTED - Creating assistant message")
                                    
                                    # Format response with source information
                                    formatted_response = response_text
                                    
                                    if sources:
                                        formatted_response += "\n\n📚 **Sources:**\n"
                                        for i, source in enumerate(sources[:3]):
                                            similarity_pct = source.get('similarity', 0) * 100
                                            formatted_response += f"{i+1}. {source.get('document', 'Unknown')} (Relevance: {similarity_pct:.1f}%)\n"
                                    
                                    # Post RAG response as assistant message
                                    response_message = self.message_post(
                                        body=formatted_response,
                                        llm_role="assistant",
                                        author_id=False
                                    )
                                    
                                    _logger.info("RAG GENERATION SUCCESSFUL")
                                    _logger.info("="*50)
                                    
                                    yield {"type": "message_create", "message": response_message.message_format()[0]}
                                    return response_message
                                else:
                                    _logger.info("RAG response not meaningful enough, falling back")
                                
                        except Exception as e:
                            _logger.error(f"RAG generation failed: {e}", exc_info=True)
                    else:
                        _logger.info("No RAG chunks available, skipping RAG")

                # Fallback to normal LLM generation
                _logger.info("FALLING BACK TO NORMAL LLM GENERATION")
                last_message = yield from self.generate_messages(last_message)
                return last_message

    def handle_append_confirmation(self, user_response):
        """Handle user's confirmation response for data appending"""
        self.ensure_one()
        
        try:
            response_lower = user_response.lower().strip()
            
            # Handle CANCEL
            if 'cancel' in response_lower:
                self._clear_classification_context()
                return self.message_post(
                    body="❌ **Operation Cancelled**\n\nData append operation has been cancelled.",
                    llm_role='assistant',
                    author_id=False
                )
            
            # Get stored context
            context = self._get_classification_context()
            if not context:
                return self.message_post(
                    body="❌ **Session expired**. Please restart the append process.",
                    llm_role='assistant',
                    author_id=False
                )
            
            _logger.info(f"DEBUG: Primary recommendation = {context['classification_result']['primary_recommendation']}")
            _logger.info(f"DEBUG: Type = {type(context['classification_result']['primary_recommendation'])}")
            
            target_table = None
            
            # Handle YES - use primary recommendation
            if response_lower == 'yes' or 'proceed' in response_lower:
                target_table = context['classification_result']['primary_recommendation']
            
            # Handle USE [table_name] - override recommendation
            elif response_lower.startswith('use '):
                table_name = response_lower.replace('use ', '').strip()
                target_table = self._resolve_table_name(table_name)
                
                if not target_table:
                    available_tables = list(self.SUPPORTED_TABLES.keys())
                    return self.message_post(
                        body=f"❌ **Unknown Table:** '{table_name}'\n\n**Available tables:** {', '.join(available_tables)}",
                        llm_role='assistant',
                        author_id=False
                    )
            
            if target_table:
                # Execute the append operation
                return self._execute_data_append(context['detected_tables'][0], target_table)
            
            # If we get here, the response wasn't recognized
            return self.message_post(
                body="❓ **Please respond with:**\n- **YES** (proceed with recommendation)\n- **USE [table_name]** (override)\n- **CANCEL** (cancel operation)",
                llm_role='assistant',
                author_id=False
            )
            
        except Exception as e:
            _logger.error(f"Error handling append confirmation: {e}", exc_info=True)
            return self.message_post(
                body=f"❌ **Error:** {str(e)}",
                llm_role='assistant',
                author_id=False
            )
            
    def append_table_data_to_odoo(self, table_data, target_table):
        """Append detected table data to specified Odoo table"""
        self.ensure_one()
        
        try:
            # Validate target table
            if target_table not in self.SUPPORTED_TABLES:
                return {
                    'success': False,
                    'message': f"Table '{target_table}' is not supported."
                }
            
            table_config = self.SUPPORTED_TABLES[target_table]
            model_name = table_config['model']
            
            # Get the target model
            target_model = self.env[model_name]
            model_fields = target_model.fields_get()
            
            # Map document columns to Odoo fields
            column_mapping = self._map_columns_to_odoo_fields(
                table_data['headers'], 
                model_fields, 
                target_table
            )
            
            if not column_mapping:
                return {
                    'success': False,
                    'message': "Could not map any document columns to Odoo fields. Please check column names.",
                    'available_fields': list(model_fields.keys())
                }
            
            # Process and validate data rows
            processed_rows = []
            errors = []
            
            # Limit rows for safety
            max_rows = min(50, len(table_data.get('data_rows', [])))
            
            for i, row in enumerate(table_data['data_rows'][:max_rows]):
                try:
                    processed_row = self._process_data_row(
                        row, column_mapping, model_fields, target_table
                    )
                    if processed_row:
                        processed_rows.append(processed_row)
                    else:
                        errors.append(f"Row {i+1}: No valid data found")
                except Exception as row_error:
                    errors.append(f"Row {i+1}: {str(row_error)}")
            
            if not processed_rows:
                return {
                    'success': False,
                    'message': f"No valid rows to insert. Errors: {'; '.join(errors[:3])}",
                    'errors': errors
                }
            
            # Create records in batches
            created_records = []
            batch_errors = []
            batch_size = 10
            
            for i in range(0, len(processed_rows), batch_size):
                batch = processed_rows[i:i+batch_size]
                try:
                    records = target_model.create(batch)
                    created_records.extend(records.ids)
                except Exception as batch_error:
                    batch_errors.append(f"Batch {i//batch_size + 1}: {str(batch_error)}")
                    # Try individual records in case of batch failure
                    for j, single_row in enumerate(batch):
                        try:
                            record = target_model.create(single_row)
                            created_records.append(record.id)
                        except Exception as single_error:
                            errors.append(f"Row {i+j+1}: {str(single_error)}")
            
            # Prepare result
            result = {
                'success': len(created_records) > 0,
                'created_count': len(created_records),
                'error_count': len(errors) + len(batch_errors),
                'created_ids': created_records,
                'errors': errors[:10],
                'batch_errors': batch_errors,
                'table_name': table_config['name'],
                'model_name': model_name,
                'column_mapping': column_mapping
            }
            
            return result
            
        except Exception as e:
            _logger.error(f"Error appending data to {target_table}: {e}", exc_info=True)
            return {
                'success': False,
                'message': f"Error appending data: {str(e)}"
            }

    def _map_columns_to_odoo_fields(self, document_headers, model_fields, target_table):
        """Map document column headers to Odoo model fields"""
        mapping = {}
        
        # Enhanced column mappings based on your specifications
        COLUMN_MAPPINGS = {
            'sale_order': {
                'name': 'name', 'customer': 'partner_id', 'partner': 'partner_id',
                'date': 'date_order', 'order_date': 'date_order', 
                'amount': 'amount_total', 'total': 'amount_total',
                'state': 'state', 'status': 'state'
            },
            'purchase_order': {
                'name': 'name', 'vendor': 'partner_id', 'supplier': 'partner_id', 
                'partner': 'partner_id', 'date': 'date_order', 'order_date': 'date_order',
                'amount': 'amount_total', 'total': 'amount_total',
                'state': 'state', 'status': 'state'
            },
            'stock_move': {
                'product': 'product_id', 'quantity': 'product_uom_qty', 'qty': 'product_uom_qty',
                'uom': 'product_uom', 'unit': 'product_uom',
                'from_location': 'location_id', 'source_location': 'location_id',
                'to_location': 'location_dest_id', 'destination_location': 'location_dest_id',
                'state': 'state', 'status': 'state'
            },
            'stock_quant': {
                'product': 'product_id', 'location': 'location_id',
                'quantity': 'quantity', 'qty': 'quantity',
                'reserved_quantity': 'reserved_quantity', 'reserved_qty': 'reserved_quantity',
                'lot': 'lot_id', 'package': 'package_id'
            },
            'stock_valuation_layer': {
                'product': 'product_id', 'quantity': 'quantity', 'qty': 'quantity',
                'value': 'value', 'unit_cost': 'unit_cost', 'cost': 'unit_cost',
                'date': 'create_date'
            },
            'hr_expense': {
                'employee': 'employee_id', 'staff': 'employee_id',
                'date': 'date', 'expense_date': 'date',
                'product': 'product_id', 'category': 'product_id',
                'amount': 'total_amount', 'total_amount': 'total_amount',
                'state': 'state', 'status': 'state'
            },
            'hr_expense_sheet': {
                'employee': 'employee_id', 'staff': 'employee_id',
                'date': 'accounting_date', 'accounting_date': 'accounting_date',
                'total_amount': 'total_amount', 'amount': 'total_amount',
                'state': 'state', 'status': 'state'
            },
            'account_move': {
                'name': 'name', 'entry_number': 'name',
                'date': 'date', 'accounting_date': 'date',
                'journal': 'journal_id', 'amount': 'amount_total',
                'total': 'amount_total', 'state': 'state', 'status': 'state'
            }
        }
        
        table_mappings = COLUMN_MAPPINGS.get(target_table, {})
        
        # Try to map each header
        for i, header in enumerate(document_headers):
            if not header:
                continue
                
            header_clean = str(header).lower().strip().replace(' ', '_').replace('-', '_')
            
            # Direct match in model fields
            if header_clean in model_fields:
                mapping[i] = header_clean
                continue
            
            # Match through predefined mappings
            matched = False
            for pattern, field_name in table_mappings.items():
                if pattern in header_clean and field_name in model_fields:
                    mapping[i] = field_name
                    matched = True
                    break
            
            # Fuzzy matching for common variations
            if not matched:
                for pattern, field_name in table_mappings.items():
                    if (pattern.replace('_', '') in header_clean.replace('_', '') or
                        header_clean.replace('_', '') in pattern.replace('_', '')):
                        if field_name in model_fields:
                            mapping[i] = field_name
                            break
        
        return mapping

    def _process_data_row(self, row, column_mapping, model_fields, target_table):
        """Process a single data row for insertion into Odoo"""
        processed_row = {}
        
        for col_index, field_name in column_mapping.items():
            if col_index >= len(row):
                continue
                
            cell_value = row[col_index]
            if not cell_value or str(cell_value).strip() == '':
                continue
                
            field_info = model_fields[field_name]
            field_type = field_info['type']
            
            try:
                # Process based on field type
                if field_type in ['char', 'text']:
                    processed_row[field_name] = str(cell_value).strip()
                    
                elif field_type == 'integer':
                    clean_value = str(cell_value).replace(',', '').replace(' ', '')
                    processed_row[field_name] = int(float(clean_value))
                    
                elif field_type in ['float', 'monetary']:
                    clean_value = re.sub(r'[^\d.-]', '', str(cell_value))
                    processed_row[field_name] = float(clean_value) if clean_value else 0.0
                    
                elif field_type == 'date':
                    processed_row[field_name] = self._parse_date(cell_value)
                    
                elif field_type == 'datetime':
                    processed_row[field_name] = self._parse_datetime(cell_value)
                    
                elif field_type == 'many2one':
                    resolved_id = self._resolve_many2one_field(cell_value, field_name, field_info)
                    if resolved_id:
                        processed_row[field_name] = resolved_id
                        
                elif field_type == 'selection':
                    selection_value = self._resolve_selection_field(cell_value, field_info)
                    if selection_value:
                        processed_row[field_name] = selection_value
                        
            except (ValueError, TypeError) as e:
                _logger.warning(f"Could not process field {field_name} with value {cell_value}: {e}")
                continue
        
        # Add required fields with defaults
        processed_row = self._add_required_defaults(processed_row, target_table)
        
        return processed_row if processed_row else None
    
    def _parse_date(self, date_value):
        """Parse various date formats"""
        if not date_value:
            return datetime.now().date()
            
        date_str = str(date_value).strip()
        
        date_formats = [
            '%Y-%m-%d', '%d/%m/%Y', '%m/%d/%Y', 
            '%d-%m-%Y', '%Y/%m/%d', '%d.%m.%Y',
            '%d %b %Y', '%d %B %Y', '%b %d, %Y', '%B %d, %Y'
        ]
        
        for fmt in date_formats:
            try:
                return datetime.strptime(date_str, fmt).date()
            except ValueError:
                continue
                
        return datetime.now().date()

    def _parse_datetime(self, datetime_value):
        """Parse datetime values"""
        if not datetime_value:
            return datetime.now()
            
        if isinstance(datetime_value, date):
            return datetime.combine(datetime_value, datetime.min.time())
            
        parsed_date = self._parse_date(datetime_value)
        return datetime.combine(parsed_date, datetime.min.time())

    def _resolve_many2one_field(self, value, field_name, field_info):
        """Resolve many2one field by searching for existing record by ID or name."""
        if not value:
            return None
            
        relation_model = self.env[field_info['relation']]
        search_value = str(value).strip()
        
        # 1. Try to match as an integer ID first. This is more reliable.
        if search_value.isdigit():
            record_id = int(search_value)
            # The browse method is very efficient for checking existence by ID
            record = relation_model.browse(record_id)
            if record.exists():
                return record.id
        
        # 2. If not found by ID, fall back to searching by name (case-insensitive)
        try:
            # Prioritize exact name match ('ilike' is case-insensitive)
            record = relation_model.search([('name', 'ilike', search_value)], limit=1)
            if record:
                return record.id

            # Broader search on display_name if available
            if 'display_name' in relation_model._fields:
                record = relation_model.search([('display_name', 'ilike', search_value)], limit=1)
                if record:
                    return record.id
            
            # For partner fields, try additional common fields
            if 'partner' in field_name and field_info['relation'] == 'res.partner':
                record = relation_model.search([
                    ('commercial_company_name', 'ilike', search_value)
                ], limit=1)
                if record:
                    return record.id
            
            # For product fields
            elif 'product' in field_name and field_info['relation'] == 'product.product':
                record = relation_model.search([
                    '|', ('default_code', '=', search_value),
                    ('barcode', '=', search_value)
                ], limit=1)
                if record:
                    return record.id
                    
        except Exception as e:
            _logger.warning(f"Could not resolve M2O field '{field_name}' with value '{search_value}': {e}")
            
        return None
    
    def _resolve_selection_field(self, value, field_info):
        """Resolve selection field values"""
        if not value:
            return None
            
        selection_options = field_info.get('selection', [])
        value_lower = str(value).lower().strip()
        
        # Try exact match first
        for option_key, option_label in selection_options:
            if option_key.lower() == value_lower:
                return option_key
            if option_label.lower() == value_lower:
                return option_key
        
        # Try partial match
        for option_key, option_label in selection_options:
            if value_lower in option_key.lower() or value_lower in option_label.lower():
                return option_key
                
        return None
    
    def _add_required_defaults(self, row_data, target_table):
        """Add required field defaults based on table configuration"""
        table_config = self.SUPPORTED_TABLES[target_table]
        default_values = table_config.get('default_values', {})
        
        # Add default values for missing fields
        for field_name, default_value in default_values.items():
            if field_name not in row_data:
                row_data[field_name] = default_value
        
        # Add specific defaults based on table type
        if target_table == 'stock_move':
            if 'name' not in row_data:
                row_data['name'] = f"Imported Stock Move - {datetime.now().strftime('%Y-%m-%d %H:%M')}"
            if 'product_uom' not in row_data:
                uom = self.env['uom.uom'].search([('name', '=', 'Units')], limit=1)
                if uom:
                    row_data['product_uom'] = uom.id
        
        elif target_table == 'account_move':
            if 'journal_id' not in row_data:
                journal = self.env['account.journal'].search([('type', '=', 'general')], limit=1)
                if journal:
                    row_data['journal_id'] = journal.id
        
        return row_data

    def _execute_data_append(self, table_data, target_table):
        """Execute the actual data append operation"""
        
        # Show processing message
        processing_msg = self.message_post(
            body=f"🔄 **Processing...** Appending data to {self.SUPPORTED_TABLES[target_table]['name']}",
            llm_role='assistant',
            author_id=False
        )
        
        # Execute the append
        result = self.append_table_data_to_odoo(table_data, target_table)
        
        # Clear context after processing
        self._clear_classification_context()
        
        # Format response based on result
        if result['success']:
            response = f"""
    ✅ **Success!** Data appended to **{result['table_name']}**

    **📊 Results:**
    - **Records Created:** {result['created_count']}
    - **Target Model:** `{result['model_name']}`
    - **Record IDs:** {', '.join(map(str, result['created_ids'][:10]))}{'...' if len(result['created_ids']) > 10 else ''}

    **🔗 Column Mapping:**
    """
            # Show which columns were mapped
            for col_idx, field_name in result.get('column_mapping', {}).items():
                header_name = table_data['headers'][col_idx] if col_idx < len(table_data.get('headers', [])) else f"Column {col_idx}"
                response += f"- `{header_name}` → `{field_name}`\n"
            
            if result.get('errors') or result.get('batch_errors'):
                total_errors = len(result.get('errors', [])) + len(result.get('batch_errors', []))
                response += f"\n⚠️ **Warnings:** {total_errors} issues encountered\n"
                for error in (result.get('errors', []) + result.get('batch_errors', []))[:3]:
                    response += f"- {error}\n"
        else:
            response = f"""
    ❌ **Failed** to append data to {self.SUPPORTED_TABLES[target_table]['name']}

    **Error:** {result.get('message', 'Unknown error')}
    """
            if result.get('errors'):
                response += f"\n**Details:**\n"
                for error in result['errors'][:5]:
                    response += f"- {error}\n"
            
            if result.get('available_fields'):
                response += f"\n**Available fields in {target_table}:** {', '.join(result['available_fields'][:10])}"
        
        return self.message_post(
            body=response,
            llm_role='assistant',
            author_id=False
        )

    def _process_user_message(self, content):
        """Process user message to detect special intents like table display, append requests, or confirmation responses"""
        _logger.info(f"Processing user message: '{content[:100]}...'")
        
        content_lower = content.lower().strip()
        
        # First check for confirmation responses (YES, NO, CANCEL, USE table_name)
        confirmation_keywords = ['yes', 'confirm', 'proceed', 'use ', 'cancel', 'no']
        if any(keyword in content_lower for keyword in confirmation_keywords):
            # Check if we have pending context
            context = self._get_classification_context()
            if context:
                _logger.info("Confirmation response detected with pending context - handling append confirmation")
                try:
                    response_message = self.handle_append_confirmation(content)
                    if response_message:
                        return response_message
                except Exception as e:
                    _logger.error(f"Confirmation handling failed: {e}", exc_info=True)
        
        # Keywords for table detection
        TABLE_DETECTION_KEYWORDS = [
            'show tables', 'show table', 'find tables', 'detect tables', 'display tables', 'list tables',
            'analyze table', 'read table', 'show data', 'what tables', 'view tables',
            'table structure', 'data structure', 'analyze document',
            'show content', 'extract tables', 'parse document', 'tables in document',
            'tabular data', 'data preview', 'spreadsheet'
        ]
        
        # Enhanced append intent detection
        EXPLICIT_APPEND_KEYWORDS = [
            'append table', 'add table', 'insert table', 'import table', 'upload table',
            'append data', 'add data', 'insert data', 'import data', 'upload data',
            'append to odoo', 'add to odoo', 'insert into odoo', 'import to odoo',
            'create records', 'add records', 'insert records'
        ]
        
        # Component-based detection
        append_verbs = ['append', 'add', 'insert', 'import', 'upload', 'create']
        data_nouns = ['table', 'data', 'records', 'information', 'entries']
        odoo_context = ['odoo', 'system', 'database', 'module']
        
        # Pattern-based check
        append_patterns = [
            r'(append|add|insert|import|upload)\s+.*?(data|table|records|information).*?(to|into)\s+.*?(odoo|system)',
            r'(append|add|insert|import|upload)\s+.*?(vendor|customer|sales|purchase|inventory|expense).*?(data|records).*?(odoo|system)',
            r'(create|add)\s+.*?(records|entries).*?(in|into|to)\s+.*?(odoo|system)'
        ]
        
        # Check for append intent
        has_explicit_append = any(keyword in content_lower for keyword in EXPLICIT_APPEND_KEYWORDS)
        has_append_verb = any(verb in content_lower for verb in append_verbs)
        has_data_noun = any(noun in content_lower for noun in data_nouns)
        has_odoo_context = any(ctx in content_lower for ctx in odoo_context)
        component_based_match = has_append_verb and has_data_noun and has_odoo_context
        pattern_based_match = any(re.search(pattern, content_lower) for pattern in append_patterns)
        
        append_intent_detected = has_explicit_append or component_based_match or pattern_based_match
        
        if append_intent_detected:
            _logger.info("Append intent detected! Processing table append request...")
            try:
                response_message = self.process_table_append_request(content)
                if response_message:
                    return response_message
                else:
                    return True
            except Exception as e:
                _logger.error(f"Append processing failed: {e}", exc_info=True)
                error_message = self.message_post(
                    body=f"Error processing append request: {str(e)}",
                    llm_role='assistant',
                    author_id=False
                )
                return error_message
        
        # Check if user wants to see tables directly
        elif any(keyword in content_lower for keyword in TABLE_DETECTION_KEYWORDS):
            _logger.info("Table display intent detected! Processing documents directly...")
            try:
                response_message = self.process_table_request(content)
                if response_message:
                    return response_message
                else:
                    return True
            except Exception as e:
                _logger.error(f"Table processing failed: {e}", exc_info=True)
                error_message = self.message_post(
                    body=f"Error processing table request: {str(e)}",
                    llm_role='assistant',
                    author_id=False
                )
                return error_message
        
        _logger.info("No special intent detected")
        return False

    # Add this method to the LLMThread model instead of using document.table.detector

    def process_table_request(self, user_message: str = ""):
        """Process user request to show tables from uploaded documents"""
        self.ensure_one()
        
        try:
            # Force refresh of attachment search to avoid caching issues
            self.env['ir.attachment'].invalidate_cache()
            self.env['mail.message'].invalidate_cache()
            
            # Get all attachments for this thread with explicit search
            thread_attachments = self.env['ir.attachment'].search([
                ('res_model', '=', 'llm.thread'),
                ('res_id', '=', self.id)
            ])
            
            # Also check message attachments - get fresh message records
            messages = self.env['mail.message'].search([
                ('model', '=', 'llm.thread'),
                ('res_id', '=', self.id)
            ])
            message_attachments = self.env['ir.attachment'].search([
                ('res_model', '=', 'mail.message'),
                ('res_id', 'in', messages.ids)
            ]) if messages else self.env['ir.attachment']
            
            # Combine all attachments
            all_attachments = thread_attachments | message_attachments
            
            _logger.info(f"Found {len(all_attachments)} total attachments for thread {self.id}")
            for att in all_attachments:
                _logger.info(f"  - {att.name} ({att.mimetype})")
            
            if not all_attachments:
                response_message = self.message_post(
                    body="No documents found in this thread. Please upload supported file types:\n"
                        "- CSV files (.csv)\n"
                        "- JSON files (.json)\n" 
                        "- Text files with tabular data (.txt)\n"
                        "- Excel files (.xlsx, .xls)\n"
                        "- Word files (.docx)",
                    llm_role='assistant',
                    author_id=False
                )
                return response_message  # ✅ Return message object
            
            # Filter for supported file types
            supported_types = [
                'text/csv', 'application/json', 'text/plain',
                'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',  # XLSX
                'application/vnd.ms-excel',  # XLS
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document'  # DOCX
            ]
            
            supported_attachments = all_attachments.filtered(
                lambda a: a.mimetype in supported_types or 
                        a.name.lower().endswith(('.csv', '.json', '.txt', '.xlsx', '.xls', '.docx'))
            )
            
            _logger.info(f"Found {len(supported_attachments)} supported attachments")
            
            if not supported_attachments:
                file_list = [f"- {att.name} ({att.mimetype})" for att in all_attachments]
                response_message = self.message_post(
                    body="No supported table files found.\n\nFound files:\n" + 
                        "\n".join(file_list) + 
                        "\n\nSupported formats: CSV, JSON, TXT, XLSX, XLS, DOCX",
                    llm_role='assistant',
                    author_id=False
                )
                return response_message  # ✅ Return message object
            
            # Process each supported attachment - ALWAYS reprocess, don't cache
            all_detected_tables = []
            for attachment in supported_attachments:
                _logger.info(f"Processing attachment: {attachment.name}")
                # Force fresh extraction each time
                tables = self._extract_tables_from_attachment(attachment)
                all_detected_tables.extend(tables)
                _logger.info(f"  - Found {len(tables)} tables")
            
            if not all_detected_tables:
                file_list = [f"- {att.name}" for att in supported_attachments]
                response_message = self.message_post(
                    body="No tables found in uploaded documents.\n\nAnalyzed files:\n" +
                        "\n".join(file_list) +
                        "\n\nMake sure your files contain clearly structured tabular data.",
                    llm_role='assistant',
                    author_id=False
                )
                return response_message  # ✅ Return message object
            
            # Always post fresh table response
            _logger.info(f"Posting response with {len(all_detected_tables)} tables")
            response_message = self._post_table_response(all_detected_tables, user_message)
            return response_message  # ✅ Return message object
            
        except Exception as e:
            _logger.error(f"Error processing table request: {e}", exc_info=True)
            error_message = self.message_post(
                body=f"Error processing documents: {str(e)}",
                llm_role='assistant',
                author_id=False
            )
            return error_message  # ✅ Return error message object

    def _extract_tables_from_attachment(self, attachment):
        """Extract tables from a single attachment"""
        tables = []
        
        try:
            # Extract text content
            text_content = self._extract_text_from_attachment(attachment)
            if not text_content:
                return tables
            
            file_ext = attachment.name.split('.')[-1].lower() if '.' in attachment.name else ''
            
            # Different extraction strategies based on file type
            if file_ext == 'csv':
                tables.extend(self._extract_csv_tables(text_content, attachment.name))
            elif file_ext == 'json':
                tables.extend(self._extract_json_tables(text_content, attachment.name))
            elif file_ext == 'docx':
                tables.extend(self._extract_docx_tables(attachment, text_content))
            elif file_ext in ['xlsx', 'xls']:
                tables.extend(self._extract_excel_tables(attachment))
            else:
                # Try generic text table detection
                tables.extend(self._extract_text_tables(text_content, attachment.name))
            
            return tables
            
        except Exception as e:
            _logger.error(f"Error extracting tables from {attachment.name}: {e}")
            return tables

    def _extract_text_from_attachment(self, attachment):
        """Extract text content from attachment"""
        try:
            if not attachment.datas:
                return None
                
            import base64
            file_data = base64.b64decode(attachment.datas)
            file_ext = attachment.name.split('.')[-1].lower() if '.' in attachment.name else ''
            
            if file_ext in ['txt', 'csv', 'json']:
                try:
                    return file_data.decode('utf-8')
                except UnicodeDecodeError:
                    return file_data.decode('latin-1', errors='ignore')
            elif file_ext == 'docx':
                return self._extract_docx_text(file_data)
            else:
                try:
                    return file_data.decode('utf-8', errors='ignore')
                except:
                    return None
                    
        except Exception as e:
            _logger.error(f"Error extracting text from {attachment.name}: {e}")
            return None

    def _extract_csv_tables(self, text_content, filename):
        """Extract CSV tables"""
        import csv
        from io import StringIO
        
        tables = []
        try:
            # Try different delimiters
            for delimiter in [',', ';', '\t', '|']:
                try:
                    csv_reader = csv.reader(StringIO(text_content), delimiter=delimiter)
                    rows = list(csv_reader)
                    
                    if len(rows) >= 2 and len(rows[0]) >= 2:
                        headers = rows[0]
                        data_rows = rows[1:]
                        
                        tables.append({
                            'name': f"CSV Table from {filename}",
                            'headers': headers,
                            'data_rows': data_rows,
                            'row_count': len(data_rows),
                            'column_count': len(headers),
                            'source': filename
                        })
                        break  # Use first successful parse
                except:
                    continue
        except Exception as e:
            _logger.error(f"CSV extraction error: {e}")
        
        return tables

    def _extract_json_tables(self, text_content, filename):
        """Extract JSON tables (arrays of objects)"""
        import json
        
        tables = []
        try:
            data = json.loads(text_content)
            
            if isinstance(data, list) and len(data) > 0 and isinstance(data[0], dict):
                headers = list(data[0].keys())
                data_rows = []
                
                for item in data:
                    if isinstance(item, dict):
                        row = [str(item.get(header, '')) for header in headers]
                        data_rows.append(row)
                
                tables.append({
                    'name': f"JSON Table from {filename}",
                    'headers': headers,
                    'data_rows': data_rows,
                    'row_count': len(data_rows),
                    'column_count': len(headers),
                    'source': filename
                })
        except Exception as e:
            _logger.error(f"JSON extraction error: {e}")
        
        return tables

    def _extract_docx_text(self, file_data):
        """Extract text from DOCX"""
        try:
            from docx import Document
            import io
            
            doc_file = io.BytesIO(file_data)
            doc = Document(doc_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    text += "\t".join(row_text) + "\n"
                text += "\n"
            
            return text.strip() if text.strip() else None
            
        except ImportError:
            _logger.warning("python-docx not available")
            return None
        except Exception as e:
            _logger.error(f"DOCX extraction error: {e}")
            return None

    def _extract_docx_tables(self, attachment, text_content):
        """Extract structured tables from DOCX"""
        tables = []
        try:
            import base64
            from docx import Document
            import io
            
            file_data = base64.b64decode(attachment.datas)
            doc_file = io.BytesIO(file_data)
            doc = Document(doc_file)
            
            for i, table in enumerate(doc.tables):
                if len(table.rows) >= 2:
                    # First row as headers
                    headers = [cell.text.strip() for cell in table.rows[0].cells]
                    
                    # Remaining rows as data
                    data_rows = []
                    for row in table.rows[1:]:
                        row_data = [cell.text.strip() for cell in row.cells]
                        data_rows.append(row_data)
                    
                    if headers and data_rows:
                        tables.append({
                            'name': f"Table {i+1} from {attachment.name}",
                            'headers': headers,
                            'data_rows': data_rows,
                            'row_count': len(data_rows),
                            'column_count': len(headers),
                            'source': attachment.name
                        })
        except Exception as e:
            _logger.error(f"DOCX table extraction error: {e}")
        
        return tables

    def _extract_text_tables(self, text_content, filename):
        """Extract tables from plain text using pattern matching"""
        import re
        
        tables = []
        try:
            lines = text_content.split('\n')
            current_table = None
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Look for tab-separated or multi-space separated values
                if re.search(r'\t|\s{3,}', line):
                    parts = re.split(r'\t|\s{3,}', line)
                    parts = [p.strip() for p in parts if p.strip()]
                    
                    if len(parts) >= 2:
                        if current_table is None:
                            current_table = {
                                'headers': parts,
                                'data_rows': [],
                                'source': filename
                            }
                        else:
                            if len(parts) == len(current_table['headers']):
                                current_table['data_rows'].append(parts)
                else:
                    if current_table and len(current_table['data_rows']) >= 1:
                        current_table.update({
                            'name': f"Text Table from {filename}",
                            'row_count': len(current_table['data_rows']),
                            'column_count': len(current_table['headers'])
                        })
                        tables.append(current_table)
                    current_table = None
            
            # Don't forget the last table
            if current_table and len(current_table['data_rows']) >= 1:
                current_table.update({
                    'name': f"Text Table from {filename}",
                    'row_count': len(current_table['data_rows']),
                    'column_count': len(current_table['headers'])
                })
                tables.append(current_table)
                
        except Exception as e:
            _logger.error(f"Text table extraction error: {e}")
        
        return tables

    def _post_table_response(self, detected_tables, user_query=""):
        """Post a response showing the detected tables using HTML format"""
        response_message = "<h3>Here are the tables from your uploaded documents:</h3><br>"
        
        for i, table in enumerate(detected_tables, 1):
            response_message += f"<h4>Table {i}: {table['name']}</h4>"
            response_message += f"<p><strong>Size:</strong> {table['row_count']} rows × {table['column_count']} columns</p>"
            
            if table['headers'] and table['data_rows']:
                # Limit columns for display
                max_columns = min(8, len(table['headers']))
                headers = table['headers'][:max_columns]
                
                # Start HTML table with styling
                response_message += '''
                <table style="border-collapse: collapse; border: 1px solid #ddd; margin: 10px 0; width: 100%; max-width: 800px;">
                <thead style="background-color: #f5f5f5;">
                <tr>'''
                
                # Add headers
                for header in headers:
                    header_text = str(header)[:30]  # Limit header length
                    response_message += f'<th style="border: 1px solid #ddd; padding: 8px; text-align: left; font-weight: bold;">{header_text}</th>'
                
                response_message += '</tr></thead><tbody>'
                
                # Add data rows (limit to 10 for performance)
                rows_to_show = min(10, len(table['data_rows']))
                for row_idx in range(rows_to_show):
                    row = table['data_rows'][row_idx]
                    response_message += '<tr>'
                    
                    for j in range(len(headers)):
                        if j < len(row):
                            cell_value = str(row[j]) if row[j] is not None else ""
                            if len(cell_value) > 50:
                                cell_value = cell_value[:47] + "..."
                            # Escape HTML characters
                            cell_value = cell_value.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        else:
                            cell_value = ""
                        
                        response_message += f'<td style="border: 1px solid #ddd; padding: 8px;">{cell_value}</td>'
                    
                    response_message += '</tr>'
                
                response_message += '</tbody></table>'
                
                if table['row_count'] > 10:
                    response_message += f"<p><em>... and {table['row_count'] - 10} more rows</em></p>"
            
            response_message += "<hr><br>"
        
        # Add summary
        total_rows = sum(table['row_count'] for table in detected_tables)
        response_message += f"<p><strong>Summary:</strong> Found {len(detected_tables)} table(s) with {total_rows} total data rows.</p>"
        
        # ✅ Return the message object instead of void
        return self.message_post(
            body=response_message,
            llm_role='assistant',
            author_id=False
        )

    # ============================================================================
    # SENTIMENT ANALYSIS AND CONFIRMATION - Enhanced Version
    # ============================================================================
    
    def classify_table_intent(self, user_message, detected_tables):
        """
        Classify user intent and determine appropriate Odoo table for data appending
        
        Args:
            user_message (str): User's message
            detected_tables (list): Tables detected from documents
            
        Returns:
            dict: Classification results
        """
        classifier = OdooTableClassifier()
        return classifier.analyze_user_intent(user_message, detected_tables)
    
    def process_table_append_request(self, user_message):
        """
        Process user request to append table data to Odoo tables with sentiment analysis
        """
        self.ensure_one()
        
        try:
            _logger.info(f"Processing table append request: '{user_message[:100]}...'")
            
            # Use the same attachment logic as process_table_request method
            all_attachments = self._get_thread_attachments()
            
            if not all_attachments:
                response_message = self.message_post(
                    body="No documents found to analyze. Please upload documents first.",
                    llm_role='assistant',
                    author_id=False
                )
                return response_message

            # Extract tables from attachments
            all_detected_tables = []
            for attachment in all_attachments:
                tables = self._extract_tables_from_attachment(attachment)
                all_detected_tables.extend(tables)
            
            if not all_detected_tables:
                response_message = self.message_post(
                    body="No tables found in the uploaded documents.",
                    llm_role='assistant',
                    author_id=False
                )
                return response_message
            
            # Perform sentiment analysis and classification
            classification_result = self.classify_table_intent(user_message, all_detected_tables)
            
            # Check if we have a confident recommendation (e.g., confidence > 40%)
            if not classification_result['primary_recommendation'] or classification_result['confidence'] < 0.4:
                response_message = self._post_unclear_intent_response(classification_result, all_detected_tables)
                return response_message
            
            self._store_classification_context(classification_result, all_detected_tables)
            # Post confirmation message
            response_message = self._post_table_append_confirmation(classification_result, all_detected_tables, user_message)
            return response_message
            
        except Exception as e:
            _logger.error(f"Error processing table append request: {e}", exc_info=True)
            error_message = self.message_post(
                body=f"Error analyzing your request: {str(e)}",
                llm_role='assistant',
                author_id=False
            )
            return error_message
    
    def _post_unclear_intent_response(self, classification_result, detected_tables):
        """Post response when intent is unclear"""
        response = "<h3>I need clarification on your intent</h3>"
        response += "<p>I detected tables in your documents but I'm not sure which Odoo module you want to append to.</p>"
        
        # Show detected tables summary
        response += f"<p><strong>Detected Tables:</strong> {len(detected_tables)} table(s)</p><ul>"
        for i, table in enumerate(detected_tables[:3]):
            response += f"<li>{table['name']} ({table['row_count']} rows)</li>"
        response += "</ul>"
        
        # Show top recommendations with corrected confidence display
        if classification_result['top_3_recommendations']:
            response += "<p><strong>Possible Odoo modules:</strong></p><ul>"
            for table, confidence in classification_result['top_3_recommendations']:
                if confidence > 0.05:  # Only show if there's at least 5% confidence
                    table_display = table.replace('_', ' ').title()
                    # Use format specifier ':.1%' to correctly format the float as a percentage
                    response += f"<li>{table_display} (Confidence: {confidence:.1%})</li>"
            response += "</ul>"
        
        response += "<p>Please clarify which type of data you want to append:</p>"
        response += "<ul>"
        response += "<li><strong>Sales Order</strong> - Customer sales, quotations, orders</li>"
        response += "<li><strong>Purchase Order</strong> - Vendor purchases, procurement</li>"
        response += "<li><strong>Inventory</strong> - Stock movements, quantities, locations</li>"
        response += "<li><strong>Expenses</strong> - Employee expenses, reimbursements</li>"
        response += "<li><strong>Accounting</strong> - Journal entries, financial records</li>"
        response += "</ul>"
        
        return self.message_post(
            body=response,
            llm_role='assistant',
            author_id=False
        )

    def _post_table_append_confirmation(self, classification_result, detected_tables, original_message):
        """Enhanced confirmation message for table append operation"""
        primary_table = classification_result['primary_recommendation']
        confidence = classification_result['confidence']
        
        # Create user-friendly table names with descriptions
        table_display_info = {
            'sale_order': {
                'name': 'Sales Orders',
                'description': 'Customer orders, quotations, and sales transactions',
                'typical_data': 'customer details, products, quantities, prices, dates'
            },
            'purchase_order': {
                'name': 'Purchase Orders', 
                'description': 'Vendor purchases and procurement orders',
                'typical_data': 'vendor details, products, quantities, costs, delivery dates'
            },
            'stock_picking': {
                'name': 'Stock Picking/Delivery Orders',
                'description': 'Warehouse operations, deliveries, and stock transfers',
                'typical_data': 'products, quantities, locations, picking dates'
            },
            'stock_move': {
                'name': 'Stock Movements',
                'description': 'Individual inventory movements and transfers',
                'typical_data': 'products, quantities, source/destination locations'
            },
            'stock_quant': {
                'name': 'Stock Quantities',
                'description': 'Current inventory levels and stock quantities',
                'typical_data': 'products, available quantities, locations, lots/serials'
            },
            'stock_validation_layer': {
                'name': 'Stock Valuation',
                'description': 'Inventory valuation and cost layers',
                'typical_data': 'products, quantities, unit costs, total values'
            },
            'hr_expense': {
                'name': 'Employee Expenses',
                'description': 'Individual employee expense entries',
                'typical_data': 'employee names, expense amounts, categories, dates'
            },
            'hr_expense_sheet': {
                'name': 'Expense Reports',
                'description': 'Employee expense reports and reimbursement requests',
                'typical_data': 'employee names, total amounts, expense periods, status'
            },
            'accounting_move': {
                'name': 'Journal Entries',
                'description': 'Accounting journal entries and financial transactions',
                'typical_data': 'accounts, debits, credits, references, posting dates'
            }
        }
        
        table_info = table_display_info.get(primary_table, {'name': primary_table, 'description': 'Unknown table'})
        
        response = f"<h3>📊 Table Classification Results</h3>"
        response += f"<p><strong>Your Request:</strong> <em>\"{original_message[:150]}{'...' if len(original_message) > 150 else ''}\"</em></p>"
        
        # Classification confidence indicator
        confidence_color = "#28a745" if confidence > 0.7 else "#ffc107" if confidence > 0.4 else "#dc3545"
        confidence_text = "High" if confidence > 0.7 else "Medium" if confidence > 0.4 else "Low"
        
        response += f"<div style='background-color: {confidence_color}; color: white; padding: 10px; border-radius: 5px; margin: 10px 0;'>"
        response += f"<h4>🎯 Recommended: {table_info['name']}</h4>"
        response += f"<p><strong>Confidence:</strong> {confidence_text} ({confidence:.1%})</p>"
        response += f"<p><strong>Description:</strong> {table_info['description']}</p>"
        response += "</div>"
        
        # Show analysis details if confidence is low
        if confidence < 0.6:
            response += "<div style='background-color: #f8f9fa; padding: 10px; border-left: 4px solid #ffc107;'>"
            response += "<h4>⚠️ Low Confidence Analysis</h4>"
            response += "<p>The classification confidence is low. Please review the recommendations carefully.</p>"
            
            # Show top alternatives
            alternatives = [item for item in classification_result['top_3_recommendations'][1:] if item[1] > 0.1]
            if alternatives:
                response += "<p><strong>Alternative options:</strong></p><ul>"
                for table, score in alternatives:
                    alt_info = table_display_info.get(table, {'name': table})
                    response += f"<li>{alt_info['name']} ({score:.1%})</li>"
                response += "</ul>"
            response += "</div>"
        
        # Document analysis summary
        total_rows = sum(table['row_count'] for table in detected_tables)
        response += f"<h4>📁 Document Analysis</h4>"
        response += f"<p><strong>Detected:</strong> {len(detected_tables)} table(s) with {total_rows} total rows</p>"
        
        # Show table preview
        response += "<div style='max-height: 200px; overflow-y: auto; border: 1px solid #ddd; padding: 10px;'>"
        for i, table in enumerate(detected_tables):
            response += f"<p><strong>Table {i+1}:</strong> {table['name']} ({table['row_count']} rows × {table['column_count']} columns)</p>"
            if table.get('headers'):
                headers_preview = ', '.join(table['headers'][:5])
                if len(table['headers']) > 5:
                    headers_preview += f" ... (+{len(table['headers'])-5} more)"
                response += f"<p><em>Columns:</em> {headers_preview}</p>"
        response += "</div>"
        
        # Clear action buttons
        response += "<div style='background-color: #e3f2fd; padding: 20px; border-radius: 5px; margin: 20px 0;'>"
        response += "<h4>🔄 Next Steps</h4>"
        response += "<p><strong>Please respond with one of the following:</strong></p>"
        response += "<div style='display: grid; grid-template-columns: 1fr 1fr; gap: 10px;'>"
        response += "<div style='background-color: #4caf50; color: white; padding: 10px; text-align: center; border-radius: 3px;'>"
        response += "<strong>\"YES\" or \"CONFIRM\"</strong><br><small>Proceed with recommended table</small>"
        response += "</div>"
        response += "<div style='background-color: #f44336; color: white; padding: 10px; text-align: center; border-radius: 3px;'>"
        response += "<strong>\"NO\" or \"CANCEL\"</strong><br><small>Cancel this operation</small>"
        response += "</div>"
        response += "</div>"
        response += f"<p><strong>Or specify:</strong> \"Use [table name]\" to override (e.g., \"Use Purchase Orders\")</p>"
        response += "</div>"
        
        # Store classification result for later confirmation handling
        _logger.info(f"Classification result for thread {self.id}: {classification_result}")
        
        # ✅ Return the message object instead of void
        return self.message_post(
            body=response,
            llm_role='assistant',
            author_id=False
        )

    def generate_messages(self, last_message):
        """Simplified generate_messages - RAG logic moved to generate() method"""
        self.ensure_one()
        
        _logger.info("GENERATE_MESSAGES CALLED (RAG logic now in generate())")
        _logger.info(f"Thread ID: {self.id}")
        
        # Get last message if not provided
        if not last_message:
            last_message = self.get_latest_llm_message()
            _logger.info("Retrieved latest LLM message")
            
        if not last_message:
            _logger.warning("No LLM messages found in thread for generation")
            return None

        _logger.info(f"Last message ID: {last_message.id}")
        _logger.info(f"Last message body preview: {last_message.body[:200]}...")
        
        # This will call the parent class implementation (likely from llm_assistant)
        # Since RAG is now handled in generate(), this is just the fallback
        _logger.info("Delegating to parent generate_messages implementation")
        try:
            return (yield from super().generate_messages(last_message))
        except Exception as e:
            _logger.error(f"Parent generate_messages failed: {e}", exc_info=True)
            
            # Create error message if parent fails
            error_message = self.message_post(
                body=f"Sorry, I encountered an error generating a response: {str(e)}",
                llm_role="assistant",
                author_id=False
            )
            yield {"type": "message_create", "message": error_message.message_format()[0]}
            return error_message

    def get_context(self, base_context=None):
        context = {
            **(base_context or {}),
            "thread_id": self.id,
        }

        try:
            related_record = self.env[self.model].browse(self.res_id)
            if related_record:
                context["related_record"] = RelatedRecordProxy(related_record)
                context["related_model"] = self.model
                context["related_res_id"] = self.res_id
            else:
                context["related_record"] = None
                context["related_model"] = None
                context["related_res_id"] = None
        except Exception as e:
            _logger.warning(
                "Error accessing related record %s,%s: %s", self.model, self.res_id, e
            )

        # Expense Analysis
        # Fixed expense analysis context - removed non-existent methods
        if self.enable_expense_analysis:
            context.update({
                'expense_analyzer': {
                    'analyze_by_product': self.analyze_expenses_by_product,
                    'analyze_trends': self.analyze_expense_trends,
                    'generate_insights': self.generate_expense_insights,
                    'get_summary': self.get_expense_summary,
                    'get_pending': self.get_pending_expenses,
                    
                }
            })
            
        # Add RAG context if enabled
        if self.rag_enabled:
            context["rag_service"] = self.env['llm.rag.service']
            context["rag_enabled"] = True
        else:
            context["rag_enabled"] = False

        return context

    # ============================================================================
    # RAG-specific methods
    # ============================================================================

    def search_documents(self, query, limit=5):
        """Search for relevant documents in this thread"""
        if not self.rag_enabled:
            return []
        
        rag_service = self.env['llm.rag.service']
        return rag_service.search_relevant_chunks(query, self.id, limit)

    def generate_rag_response(self, query, context_limit=3):
        """Generate a response using RAG"""
        if not self.rag_enabled:
            raise UserError(_("RAG is not enabled for this thread."))
        
        rag_service = self.env['llm.rag.service']
        return rag_service.generate_rag_response(query, self.id, context_limit)
    
    # ============================================================================
    # EXPENSE ANALYSIS METHODS - Custom Finance AI Integration
    # ============================================================================

    # Fixed expense analysis methods for LLM thread

    def _is_expense_query(self, message):
        expense_keywords = [
            'expense', 'expenses', 'spending', 'spend', 'cost', 'budget',
            'financial', 'money', 'expense analysis', 'analyze expenses',
            'spend analysis', 'by product', 'insights', 'expenditure',
            'breakdown', 'pending', 'approval', 'reimburse'
        ]
        m = (message or "").lower()
        return any(k in m for k in expense_keywords)

    def _extract_timeframe(self, message):
        m = (message or "").lower()
        if 'this quarter' in m or 'current quarter' in m:
            return 'this_quarter'
        elif 'this month' in m or 'current month' in m:
            return 'this_month'
        elif 'last month' in m or 'previous month' in m:
            return 'last_month'
        elif 'this year' in m or 'current year' in m:
            return 'this_year'
        elif 'last 30 days' in m or '30 days' in m:
            return 'last_30_days'
        return 'this_month'

    def _parse_month_year(self, text):
        try:
            import calendar, re, datetime
            m = re.search(r'(january|february|march|april|may|june|july|august|september|october|november|december)\s+(\d{4})', (text or '').lower())
            if not m:
                return None, None
            month = list(calendar.month_name).index(m.group(1).capitalize())
            year = int(m.group(2))
            start = datetime.date(year, month, 1)
            last_day = calendar.monthrange(year, month)[1]
            end = datetime.date(year, month, last_day)
            return start, end
        except Exception as e:
            _logger.warning("Error parsing month/year: %s", e)
            return None, None

    def generate_expense_response(self, message, **kwargs):
        self.ensure_one()
        m = (message or "").lower()
        try:
            # No trend route anymore
            if 'pending' in m or 'approval' in m or 'reimburse' in m:
                resp = self.get_pending_expenses()
            elif 'summary' in m or 'overview' in m:
                resp = self.get_expense_summary()
            elif 'insight' in m or 'recommendation' in m or 'ai' in m:
                resp = self.generate_expense_insights(message)
            else:
                # Default: by-product with month-year support or fixed timeframe
                start, end = self._parse_month_year(message)
                if start and end:
                    resp = self.analyze_expenses_by_product(date_from=start, date_to=end)
                else:
                    timeframe = self._extract_timeframe(message)
                    resp = self.analyze_expenses_by_product(timeframe)
            if not resp or resp.strip() == "":
                resp = "No expense data found or unable to generate a response."
            return resp
        except Exception as e:
            self.env.cr.rollback()
            _logger.exception("LLMThread(%s): error in generate_expense_response: %s", self.id, e)
            return f"❌ Error generating expense response: {str(e)}"

    def _extract_product_name(self, message):
        """Extract product name from message."""
        try:
            # Split message into tokens and look for capitalized words
            tokens = [w.strip(",.!?:;") for w in message.split()]
            candidates = [w for w in tokens if w.istitle() and len(w) > 2]
            
            if candidates:
                # Join consecutive capitalized words (e.g., "Laptop Computer")
                product_name = ' '.join(candidates[:2])  # Take first 2 words max
                return product_name
            
            # Fallback: search for products in database
            ProductTmpl = self.env['product.template']
            for word in tokens:
                if len(word) > 3:  # Only check meaningful words
                    rec = ProductTmpl.search([('name', 'ilike', word)], limit=1)
                    if rec:
                        return rec.name
            
            return None
        except Exception as e:
            _logger.warning(f"Error extracting product name: {e}")
            return None

    def analyze_expenses_by_product(self, timeframe='this_month', date_from=None, date_to=None):
        try:
            import calendar, re, datetime
            today = datetime.date.today()
            if date_from and date_to:
                start, end = date_from, date_to
            else:
                end = today 
                if timeframe == 'this_quarter':
                    month = (today.month - 1) // 3 * 3 + 1
                    start = datetime.date(today.year, month, 1)
                elif timeframe == 'this_year':
                    start = datetime.date(today.year, 1, 1)
                elif timeframe == 'last_month':
                    import calendar
                    if today.month == 1:
                        # last month is December of previous year
                        start = datetime.date(today.year - 1, 12, 1)
                        end = datetime.date(today.year - 1, 12, calendar.monthrange(today.year - 1, 12)[1])
                    else:
                        start = datetime.date(today.year, today.month - 1, 1)
                        end = datetime.date(today.year, today.month - 1, calendar.monthrange(today.year, today.month - 1)[1])
                elif timeframe == 'last_30_days':
                    start = end - datetime.timedelta(days=30)
                else:
                    start = end.replace(day=1)

            domain = [
                ('state', 'in', ['approved']),
                ('date', '>=', start),
                ('date', '<=', end),
                ('product_id', '!=', False),
            ]
            
            expenses = self.env['hr.expense'].search(domain, order='date desc')
            
            if not expenses:
                return f'<div style="background: #d4edda; padding: 15px; border-radius: 8px; color: #155724; margin: 10px 0; text-align: center;"><h4>📊 No Expense Data Found</h4><p>No approved expenses found for the period {start.strftime("%b %d, %Y")} to {end.strftime("%b %d, %Y")}</p></div>'
            
            # Group by product for summary table
            product_totals = {}
            expense_details = []
            
            total_amount = 0.0
            for e in expenses:
                emp = e.employee_id.name or 'Unknown Employee'
                dt = e.date.strftime("%Y-%m-%d") if e.date else "No Date"
                amt = e.total_amount or 0.0
                product_name = e.product_id.product_tmpl_id.name or e.product_id.display_name or e.name or 'Other Expenses'
                
                total_amount += amt
                
                # Group by product for summary
                if product_name in product_totals:
                    product_totals[product_name]['amount'] += amt
                    product_totals[product_name]['count'] += 1
                else:
                    product_totals[product_name] = {'amount': amt, 'count': 1}
                
                # Store individual expense details
                expense_details.append({
                    'name': e.name,
                    'employee': emp,
                    'date': dt,
                    'amount': amt,
                    'product': product_name,
                    'state': e.state
                })
            
            # Build HTML output
            html_lines = []
            
            # Header with period info
            html_lines.append('<div style="background: linear-gradient(135deg, #28a745, #20c997); color: white; padding: 15px; border-radius: 8px; margin: 10px 0; text-align: center;">')
            html_lines.append('<h2 style="margin: 0 0 5px 0; font-size: 18px;">📊 EXPENSE ANALYSIS BY PRODUCT</h2>')
            html_lines.append(f'<p style="margin: 0; font-size: 13px; opacity: 0.9;">📅 Period: {start.strftime("%b %d, %Y")} to {end.strftime("%b %d, %Y")}</p>')
            html_lines.append(f'<p style="margin: 0; font-size: 13px; opacity: 0.9;">💰 Total Analyzed: ₹{total_amount:,.2f} ({len(expenses)} expenses)</p>')
            html_lines.append('</div>')
            
            # Product summary table
            if product_totals:
                html_lines.append('<div style="margin: 15px 0; border-radius: 6px; overflow: hidden; border: 1px solid #ddd;">')
                html_lines.append('<h4 style="color: white; background: #28a745; padding: 8px; margin: 0; font-size: 14px;">📈 EXPENSES BY PRODUCT/SERVICE</h4>')
                
                html_lines.append('<table style="width: 100%; border-collapse: collapse; font-size: 13px;">')
                html_lines.append('<thead><tr style="background-color: #f5f5f5;">')
                html_lines.append('<th style="padding: 8px; text-align: left; font-size: 12px;">Product/Service</th>')
                html_lines.append('<th style="padding: 8px; text-align: right; font-size: 12px;">Amount</th>')
                html_lines.append('<th style="padding: 8px; text-align: right; font-size: 12px;">Share</th>')
                html_lines.append('</tr></thead><tbody>')
                
                # Sort by amount descending
                sorted_products = sorted(product_totals.items(), key=lambda x: x[1]['amount'], reverse=True)
                
                for i, (product, data) in enumerate(sorted_products):
                    pct = (data['amount'] / total_amount) * 100 if total_amount > 0 else 0
                    name = (product[:30] + '...') if len(product) > 30 else product
                    row_bg = '#f9f9f9' if i % 2 == 0 else 'white'
                    
                    html_lines.append(f'<tr style="background-color: {row_bg};">')
                    html_lines.append(f'<td style="padding: 6px; color: #333;" title="{product}">{name}</td>')
                    html_lines.append(f'<td style="padding: 6px; text-align: right; font-weight: bold; color: #28a745;">₹{data["amount"]:,.2f}</td>')
                    html_lines.append(f'<td style="padding: 6px; text-align: right; color: #28a745; font-weight: bold;">{pct:.1f}%</td>')
                    html_lines.append('</tr>')
                
                # Total row
                html_lines.append('<tr style="background: #e8f5e8; font-weight: bold; border-top: 2px solid #28a745;">')
                html_lines.append('<td style="padding: 8px; color: #155724;">TOTAL</td>')
                html_lines.append(f'<td style="padding: 8px; text-align: right; color: #155724;">₹{total_amount:,.2f}</td>')
                html_lines.append('<td style="padding: 8px; text-align: right; color: #155724;">100.0%</td>')
                html_lines.append('</tr>')
                
                html_lines.append('</tbody></table></div>')
            
            # Detailed expense list table
            html_lines.append('<div style="margin: 15px 0; border-radius: 6px; overflow: hidden; border: 1px solid #ddd;">')
            html_lines.append('<h4 style="color: white; background: #6c757d; padding: 8px; margin: 0; font-size: 14px;">📋 DETAILED EXPENSE BREAKDOWN</h4>')
            
            html_lines.append('<table style="width: 100%; border-collapse: collapse; font-size: 12px;">')
            html_lines.append('<thead><tr style="background-color: #f5f5f5;">')
            html_lines.append('<th style="padding: 6px; text-align: left; font-size: 11px;">Description</th>')
            html_lines.append('<th style="padding: 6px; text-align: left; font-size: 11px;">Employee</th>')
            html_lines.append('<th style="padding: 6px; text-align: center; font-size: 11px;">Date</th>')
            html_lines.append('<th style="padding: 6px; text-align: right; font-size: 11px;">Amount</th>')
            html_lines.append('<th style="padding: 6px; text-align: center; font-size: 11px;">Status</th>')
            html_lines.append('</tr></thead><tbody>')
            
            # Show top 15 individual expenses
            for i, detail in enumerate(expense_details[:15]):
                desc = (detail['name'][:25] + '...') if len(detail['name']) > 25 else detail['name']
                emp = (detail['employee'][:15] + '...') if len(detail['employee']) > 15 else detail['employee']
                row_bg = '#f9f9f9' if i % 2 == 0 else 'white'
                
                html_lines.append(f'<tr style="background-color: {row_bg};">')
                html_lines.append(f'<td style="padding: 5px; color: #333;" title="{detail["name"]}">{desc}</td>')
                html_lines.append(f'<td style="padding: 5px; color: #333;" title="{detail["employee"]}">{emp}</td>')
                html_lines.append(f'<td style="padding: 5px; text-align: center; color: #666;">{detail["date"]}</td>')
                html_lines.append(f'<td style="padding: 5px; text-align: right; font-weight: bold; color: #28a745;">₹{detail["amount"]:,.2f}</td>')
                html_lines.append(f'<td style="padding: 5px; text-align: center; color: #17a2b8; font-weight: bold;">{detail["state"]}</td>')
                html_lines.append('</tr>')
            
            if len(expense_details) > 15:
                remaining_count = len(expense_details) - 15
                remaining_amount = sum(d['amount'] for d in expense_details[15:])
                html_lines.append('<tr style="background: #fff3cd; font-style: italic;">')
                html_lines.append(f'<td colspan="4" style="padding: 8px; color: #856404;">... and {remaining_count} more expenses totaling ₹{remaining_amount:,.2f}</td>')
                html_lines.append('<td style="padding: 8px; text-align: center; color: #856404;">approved</td>')
                html_lines.append('</tr>')
            
            # Final total row
            html_lines.append('<tr style="background: #e8f5e8; font-weight: bold; border-top: 2px solid #28a745;">')
            html_lines.append('<td colspan="3" style="padding: 8px; color: #155724;">TOTAL ANALYZED</td>')
            html_lines.append(f'<td style="padding: 8px; text-align: right; color: #155724;">₹{total_amount:,.2f}</td>')
            html_lines.append('<td style="padding: 8px; text-align: center; color: #155724;">-</td>')
            html_lines.append('</tr>')
            
            html_lines.append('</tbody></table></div>')
            
            return '\n'.join(html_lines)
            
        except Exception as e:
            self.env.cr.rollback()
            _logger.exception("Error in analyze_expenses_by_product: %s", e)
            return f'<div style="background: #f8d7da; padding: 15px; border-radius: 8px; color: #721c24; margin: 10px 0;"><h4>❌ Error</h4><p>Error analyzing expenses: {str(e)}</p></div>'

    def get_pending_expenses(self):
        try:
            pending_states = ['reported']  # pending == reported (to be reimbursed)
            expenses = self.env['hr.expense'].search(
                [('state', 'in', pending_states)],
                order='date desc', limit=50
            )
            if not expenses:
                return '<div style="background: #d4edda; padding: 15px; border-radius: 8px; color: #155724; margin: 10px 0; text-align: center;"><h4>✅ No Pending Expenses</h4><p>All expenses are processed!</p></div>'
            
            # Group by product for summary table
            product_totals = {}
            employee_details = []
            
            total_pending = 0.0
            for e in expenses:
                emp = e.employee_id.name or 'Unknown Employee'
                dt = e.date.strftime("%Y-%m-%d") if e.date else "No Date"
                amt = e.total_amount or 0.0
                product_name = e.product_id.name or e.name or 'Other Expenses'
                
                total_pending += amt
                
                # Group by product for summary
                if product_name in product_totals:
                    product_totals[product_name] += amt
                else:
                    product_totals[product_name] = amt
                
                # Store individual expense details
                employee_details.append({
                    'name': e.name,
                    'employee': emp,
                    'date': dt,
                    'amount': amt,
                    'product': product_name,
                    'state': e.state
                })
            
            # Build HTML output matching the expense insights format
            html_lines = []
            
            # Header
            html_lines.append('<div style="background: linear-gradient(135deg, #dc3545, #fd7e14); color: white; padding: 15px; border-radius: 8px; margin: 10px 0; text-align: center;">')
            html_lines.append('<h2 style="margin: 0 0 5px 0; font-size: 18px;">⏳ PENDING EXPENSE REIMBURSEMENTS</h2>')
            html_lines.append(f'<p style="margin: 0; font-size: 13px; opacity: 0.9;">💰 Total Pending: ₹{total_pending:,.2f} ({len(expenses)} expenses)</p>')
            html_lines.append('</div>')
            
            # Product summary table (using same format as expense insights)
            if product_totals:
                html_lines.append('<div style="margin: 15px 0; border-radius: 6px; overflow: hidden; border: 1px solid #ddd;">')
                html_lines.append('<h4 style="color: white; background: #dc3545; padding: 8px; margin: 0; font-size: 14px;">📊 PENDING EXPENSES BY PRODUCT/SERVICE</h4>')
                
                html_lines.append('<table style="width: 100%; border-collapse: collapse; font-size: 13px;">')
                html_lines.append('<thead><tr style="background-color: #f5f5f5;">')
                html_lines.append('<th style="padding: 8px; text-align: left; font-size: 12px;">Product/Service</th>')
                html_lines.append('<th style="padding: 8px; text-align: right; font-size: 12px;">Amount</th>')
                html_lines.append('<th style="padding: 8px; text-align: right; font-size: 12px;">Share</th>')
                html_lines.append('</tr></thead><tbody>')
                
                # Sort by amount descending
                sorted_products = sorted(product_totals.items(), key=lambda x: x[1], reverse=True)
                
                for i, (product, amt) in enumerate(sorted_products[:8]):  # Show top 8
                    pct = (amt / total_pending) * 100 if total_pending > 0 else 0
                    name = (product[:30] + '...') if len(product) > 30 else product
                    row_bg = '#f9f9f9' if i % 2 == 0 else 'white'
                    
                    html_lines.append(f'<tr style="background-color: {row_bg};">')
                    html_lines.append(f'<td style="padding: 6px; color: #333;" title="{product}">{name}</td>')
                    html_lines.append(f'<td style="padding: 6px; text-align: right; font-weight: bold; color: #dc3545;">₹{amt:,.2f}</td>')
                    html_lines.append(f'<td style="padding: 6px; text-align: right; color: #dc3545; font-weight: bold;">{pct:.1f}%</td>')
                    html_lines.append('</tr>')
                
                # Total row
                html_lines.append('<tr style="background: #ffeaa7; font-weight: bold; border-top: 2px solid #fdcb6e;">')
                html_lines.append('<td style="padding: 8px; color: #e17055;">TOTAL</td>')
                html_lines.append(f'<td style="padding: 8px; text-align: right; color: #e17055;">₹{total_pending:,.2f}</td>')
                html_lines.append('<td style="padding: 8px; text-align: right; color: #e17055;">100.0%</td>')
                html_lines.append('</tr>')
                
                html_lines.append('</tbody></table></div>')
            
            # Detailed expense list table
            html_lines.append('<div style="margin: 15px 0; border-radius: 6px; overflow: hidden; border: 1px solid #ddd;">')
            html_lines.append('<h4 style="color: white; background: #6c757d; padding: 8px; margin: 0; font-size: 14px;">📋 DETAILED PENDING EXPENSES</h4>')
            
            html_lines.append('<table style="width: 100%; border-collapse: collapse; font-size: 12px;">')
            html_lines.append('<thead><tr style="background-color: #f5f5f5;">')
            html_lines.append('<th style="padding: 6px; text-align: left; font-size: 11px;">Description</th>')
            html_lines.append('<th style="padding: 6px; text-align: left; font-size: 11px;">Employee</th>')
            html_lines.append('<th style="padding: 6px; text-align: center; font-size: 11px;">Date</th>')
            html_lines.append('<th style="padding: 6px; text-align: right; font-size: 11px;">Amount</th>')
            html_lines.append('<th style="padding: 6px; text-align: center; font-size: 11px;">Status</th>')
            html_lines.append('</tr></thead><tbody>')
            
            # Show top 15 individual expenses
            for i, detail in enumerate(employee_details[:15]):
                desc = (detail['name'][:25] + '...') if len(detail['name']) > 25 else detail['name']
                emp = (detail['employee'][:15] + '...') if len(detail['employee']) > 15 else detail['employee']
                row_bg = '#f9f9f9' if i % 2 == 0 else 'white'
                
                html_lines.append(f'<tr style="background-color: {row_bg};">')
                html_lines.append(f'<td style="padding: 5px; color: #333;" title="{detail["name"]}">{desc}</td>')
                html_lines.append(f'<td style="padding: 5px; color: #333;" title="{detail["employee"]}">{emp}</td>')
                html_lines.append(f'<td style="padding: 5px; text-align: center; color: #666;">{detail["date"]}</td>')
                html_lines.append(f'<td style="padding: 5px; text-align: right; font-weight: bold; color: #dc3545;">₹{detail["amount"]:,.2f}</td>')
                html_lines.append(f'<td style="padding: 5px; text-align: center; color: #fd7e14; font-weight: bold;">{detail["state"]}</td>')
                html_lines.append('</tr>')
            
            if len(employee_details) > 15:
                remaining_count = len(employee_details) - 15
                remaining_amount = sum(d['amount'] for d in employee_details[15:])
                html_lines.append('<tr style="background: #fff3cd; font-style: italic;">')
                html_lines.append(f'<td colspan="4" style="padding: 8px; color: #856404;">... and {remaining_count} more expenses totaling ₹{remaining_amount:,.2f}</td>')
                html_lines.append('<td style="padding: 8px; text-align: center; color: #856404;">pending</td>')
                html_lines.append('</tr>')
            
            # Final total row
            html_lines.append('<tr style="background: #ffeaa7; font-weight: bold; border-top: 2px solid #fdcb6e;">')
            html_lines.append('<td colspan="3" style="padding: 8px; color: #e17055;">TOTAL PENDING</td>')
            html_lines.append(f'<td style="padding: 8px; text-align: right; color: #e17055;">₹{total_pending:,.2f}</td>')
            html_lines.append('<td style="padding: 8px; text-align: center; color: #e17055;">-</td>')
            html_lines.append('</tr>')
            
            html_lines.append('</tbody></table></div>')
            
            return '\n'.join(html_lines)
            
        except Exception as e:
            self.env.cr.rollback()
            _logger.exception("Error getting pending expenses: %s", e)
            return f'<div style="background: #f8d7da; padding: 15px; border-radius: 8px; color: #721c24; margin: 10px 0;"><h4>❌ Error</h4><p>{str(e)}</p></div>'

    def _ascii_bars(self, totals):
        """HTML bar chart with properly colored and sized progress bars"""
        if not totals:
            return '<div style="text-align: center; color: #666; padding: 20px;">No data available</div>'
        
        total_amt = sum(a for _, a in totals) or 1.0
        html_lines = ['<div style="font-family: Arial, sans-serif; line-height: 1.6; margin: 10px 0; background: #f8f9fa; padding: 15px; border-radius: 6px;">']
        max_bar_width = 300  # Fixed container width
    
        for label, amt in totals:
            pct = amt / total_amt
            # Calculate proportional width - this is key!
            bar_width = max(int(pct * max_bar_width), 10)  # Minimum 10px for visibility
            label_disp = label if len(label) <= 25 else label[:22] + "..."
            
            html_lines.append('<div style="display: flex; align-items: center; margin: 8px 0; padding: 3px 0;">')
            
            # Label with fixed width
            html_lines.append(f'<div title="{label}" style="width: 180px; font-weight: bold; color: #333; margin-right: 12px; font-size: 14px; white-space: nowrap; overflow: hidden; text-overflow: ellipsis;">{label_disp}</div>')
            
            # Fixed container (gray track) - SAME SIZE FOR ALL
            html_lines.append(f'<div style="width: {max_bar_width}px; height: 20px; background-color: #e8e8e8; border-radius: 10px; margin-right: 12px; border: 1px solid #ccc; position: relative; overflow: hidden;">')
            
            # Variable colored bar (green fill) - SIZE VARIES BY PERCENTAGE
            html_lines.append(f'<div style="height: 100%; width: {bar_width}px; background: linear-gradient(90deg, #4caf50, #66bb6a); border-radius: 9px; box-shadow: inset 0 1px 2px rgba(255,255,255,0.3); transition: width 0.3s ease;"></div>')
            
            html_lines.append('</div>') # Close bar container
            
            # Percentage with better styling
            html_lines.append(f'<div style="width: 60px; text-align: right; font-weight: bold; color: #2e7d32; font-size: 13px;">{pct*100:.1f}%</div>')
            
            html_lines.append('</div>') # Close row container
        
        html_lines.append('</div>')
        return '\n'.join(html_lines)

    def _totals_by_product(self, state_list, start, end):
        """No changes - keeping original logic"""
        exps = self.env['hr.expense'].search([
            ('state','in', state_list),
            ('date','>=', start), ('date','<=', end),
            ('product_id','!=', False),
        ])
        totals = {}
        for e in exps:
            pname = e.product_id.product_tmpl_id.name or e.product_id.display_name or e.name or 'Unknown'
            totals[pname] = totals.get(pname, 0.0) + (e.total_amount or 0.0)
        return sorted(totals.items(), key=lambda kv: kv[1], reverse=True)

    def _format_currency(self, amt):
        """No changes - keeping original logic"""
        return f"₹{amt:,.2f}"

    def _simple_table(self, totals, title):
        """Compact HTML table"""
        if not totals:
            return f'<div style="margin: 10px 0; padding: 10px; background: #fff3e0; border-radius: 6px; color: #666;"><strong>{title}</strong><br>No data available</div>'
        
        total_amt = sum(a for _, a in totals) or 1.0
        html_lines = []
        html_lines.append('<div style="margin: 15px 0; border-radius: 6px; overflow: hidden; border: 1px solid #ddd;">')
        html_lines.append(f'<h4 style="color: white; background: #1976d2; padding: 8px; margin: 0; font-size: 14px;">{title}</h4>')
        
        html_lines.append('<table style="width: 100%; border-collapse: collapse; font-size: 13px;">')
        html_lines.append('<thead><tr style="background-color: #f5f5f5;">')
        html_lines.append('<th style="padding: 8px; text-align: left; font-size: 12px;">Product/Service</th>')
        html_lines.append('<th style="padding: 8px; text-align: right; font-size: 12px;">Amount</th>')
        html_lines.append('<th style="padding: 8px; text-align: right; font-size: 12px;">Share</th>')
        html_lines.append('</tr></thead><tbody>')
        
        for i, (label, amt) in enumerate(totals[:8]):  # Show only top 8
            pct = (amt / total_amt) * 100
            name = (label[:30] + '...') if len(label) > 30 else label
            row_bg = '#f9f9f9' if i % 2 == 0 else 'white'
            
            html_lines.append(f'<tr style="background-color: {row_bg};">')
            html_lines.append(f'<td style="padding: 6px; color: #333;" title="{label}">{name}</td>')
            html_lines.append(f'<td style="padding: 6px; text-align: right; font-weight: bold; color: #2e7d32;">{self._format_currency(amt)}</td>')
            html_lines.append(f'<td style="padding: 6px; text-align: right; color: #1976d2; font-weight: bold;">{pct:.1f}%</td>')
            html_lines.append('</tr>')
        
        # Compact total row
        html_lines.append('<tr style="background: #e8f5e9; font-weight: bold; border-top: 2px solid #4caf50;">')
        html_lines.append('<td style="padding: 8px; color: #2e7d32;">TOTAL</td>')
        html_lines.append(f'<td style="padding: 8px; text-align: right; color: #2e7d32;">{self._format_currency(total_amt)}</td>')
        html_lines.append('<td style="padding: 8px; text-align: right; color: #2e7d32;">100.0%</td>')
        html_lines.append('</tr>')
        
        html_lines.append('</tbody></table></div>')
        return '\n'.join(html_lines)

    def _summary_stats(self, approved_totals, reported_totals, all_totals):
        """Compact summary statistics"""
        ap = sum(a for _, a in approved_totals) if approved_totals else 0.0
        rp = sum(a for _, a in reported_totals) if reported_totals else 0.0
        al = sum(a for _, a in all_totals) if all_totals else 0.0
        top_name, top_amt = ("None", 0.0)
        if all_totals:
            top_name, top_amt = all_totals[0]
        
        html_lines = []
        html_lines.append('<div style="background: #f3e5f5; padding: 15px; border-radius: 8px; margin: 15px 0; border-left: 4px solid #9c27b0;">')
        html_lines.append('<h3 style="color: #6a1b9a; margin: 0 0 10px 0; font-size: 16px; text-align: center;">SUMMARY STATISTICS</h3>')
        
        # Horizontal stats layout
        html_lines.append('<div style="display: flex; gap: 15px; margin-bottom: 10px; flex-wrap: wrap;">')
        
        # Three stat boxes in a row
        stats = [
            ("✅", "Approved", ap, "#4caf50"),
            ("📋", "Reported", rp, "#ff9800"), 
            ("📈", "All States", al, "#2196f3")
        ]
        
        for icon, label, amount, color in stats:
            html_lines.append(f'<div style="background: white; padding: 10px; border-radius: 6px; flex: 1; min-width: 120px; border-left: 3px solid {color};">')
            html_lines.append(f'<div style="font-size: 12px; color: {color}; font-weight: bold;">{icon} {label}</div>')
            html_lines.append(f'<div style="color: {color}; font-size: 16px; font-weight: bold;">{self._format_currency(amount)}</div>')
            html_lines.append('</div>')
        
        html_lines.append('</div>')
        
        # Top expense in a compact format
        if al > 0 and top_amt > 0:
            share = (top_amt / al) * 100
            html_lines.append('<div style="background: white; padding: 10px; border-radius: 6px; border-left: 3px solid #e91e63;">')
            html_lines.append(f'<div style="font-size: 12px; color: #c2185b; font-weight: bold;">🔝 Top Category: {top_name}</div>')
            html_lines.append(f'<div style="font-size: 14px; color: #c2185b; font-weight: bold;">{self._format_currency(top_amt)} ({share:.1f}%)</div>')
            html_lines.append('</div>')
        
        html_lines.append('</div>')
        return '\n'.join(html_lines)

    def _visual_chart(self, totals, title):
        """Compact visual chart with proper header"""
        html_lines = []
        html_lines.append('<div style="background: #1976d2; color: white; padding: 8px; border-radius: 6px 6px 0 0; margin: 15px 0 0 0;">')
        html_lines.append(f'<h4 style="margin: 0; font-size: 14px; text-align: center;">{title}</h4>')
        html_lines.append('</div>')
        
        if not totals:
            html_lines.append('<div style="background: white; padding: 15px; text-align: center; color: #666; border-radius: 0 0 6px 6px; border: 1px solid #ddd; border-top: none;">No data available</div>')
            return '\n'.join(html_lines)
        
        html_lines.append('<div style="background: white; border-radius: 0 0 6px 6px; border: 1px solid #ddd; border-top: none;">')
        html_lines.append(self._ascii_bars(totals))
        html_lines.append('</div>')
        
        return '\n'.join(html_lines)

    def generate_expense_insights(self, prompt):
        """Fixed expense insights with proper header and bar colors"""
        try:
            import datetime
            # Resolve period (month-year preferred if present)
            start, end = self._parse_month_year(prompt)
            if not (start and end):
                today = datetime.date.today()
                month = (today.month - 1) // 3 * 3 + 1
                start = datetime.date(today.year, month, 1)
                end = today
            
            # Build three visualizations by product
            approved_totals = self._totals_by_product(['approved'], start, end)
            reported_totals = self._totals_by_product(['reported'], start, end)
            all_totals = self._totals_by_product(['approved','reported','to_submit'], start, end)
            
            # Fixed header with proper styling
            period_header = '<div style="background: linear-gradient(135deg, #2c3e50, #3498db); color: white; padding: 20px; border-radius: 8px; margin: 10px 0; text-align: center; box-shadow: 0 4px 6px rgba(0,0,0,0.1);">'
            period_header += '<h2 style="margin: 0 0 8px 0; font-size: 22px; font-weight: bold;">EXPENSE ANALYSIS REPORT</h2>'
            period_header += f'<p style="margin: 0; font-size: 14px; opacity: 0.9; font-weight: 300;">{start.strftime("%B %d, %Y")} to {end.strftime("%B %d, %Y")}</p>'
            period_header += '</div>'
            
            # Generate all sections with reduced spacing
            summary_stats = self._summary_stats(approved_totals, reported_totals, all_totals)
            approved_table = self._simple_table(approved_totals, "✅ APPROVED EXPENSES")
            approved_chart = self._visual_chart(approved_totals, "APPROVED BREAKDOWN")
            all_table = self._simple_table(all_totals, "📊 ALL EXPENSES")
            all_chart = self._visual_chart(all_totals, "ALL EXPENSES BREAKDOWN")
            
            # Include breakdown text from canonical function
            try:
                breakdown = self.analyze_expenses_by_product(date_from=start, date_to=end)
            except:
                breakdown = "Detailed breakdown unavailable"
            
            # Combine with minimal spacing
            viz = period_header + summary_stats + approved_table + approved_chart + all_table + all_chart
            
            # Compact AI insights
            try:
                ai_prompt = f"""Based on expense data, provide 3 brief insights:
                - Approved: ₹{sum(a for _, a in approved_totals) if approved_totals else 0:,.2f}
                - Total: ₹{sum(a for _, a in all_totals) if all_totals else 0:,.2f}
                - Top: {'; '.join([f"{name}: ₹{amt:,.2f}" for name, amt in all_totals[:3]]) if all_totals else 'None'}
                User: {prompt}"""
                
                import requests
                resp = requests.post(
                    f'{self.ollama_endpoint}/api/generate',
                    json={
                        'model': self.ollama_model, 
                        'prompt': ai_prompt, 
                        'stream': False,
                        'options': {'temperature': 0.5, 'max_tokens': 300}
                    },
                    timeout=30
                )
                
                if resp.status_code == 200:
                    ai_response = resp.json().get('response', '')
                    ai_insights = '<div style="background: #8e24aa; color: white; padding: 12px; border-radius: 8px; margin: 15px 0;">'
                    ai_insights += '<h4 style="margin: 0 0 8px 0; font-size: 14px;">🤖 AI INSIGHTS</h4>'
                    ai_insights += f'<div style="font-size: 13px; line-height: 1.4;">{ai_response}</div></div>'
                    return viz + ai_insights
                else:
                    return viz + '<div style="background: #fff3cd; padding: 8px; border-radius: 6px; color: #856404; margin: 10px 0; font-size: 12px;">⚠ AI analysis unavailable</div>'
                    
            except ImportError:
                return viz + '<div style="background: #fff3cd; padding: 8px; border-radius: 6px; color: #856404; margin: 10px 0; font-size: 12px;">⚠ AI analysis unavailable (requests library missing)</div>'
            except Exception as e:
                return viz + '<div style="background: #f8d7da; padding: 8px; border-radius: 6px; color: #721c24; margin: 10px 0; font-size: 12px;">❌ AI analysis error</div>'
                
        except Exception as e:
            self.env.cr.rollback()
            _logger.exception("Error generating expense insights: %s", e)
            return f'<div style="background: #f8d7da; padding: 15px; border-radius: 8px; color: #721c24; margin: 10px 0;"><h4>❌ Error</h4><p>{str(e)}</p></div>'

    # ============================================================================
    # POSTGRESQL ADVISORY LOCK IMPLEMENTATION
    # ============================================================================

    def _acquire_thread_lock(self):
        """Acquire PostgreSQL advisory lock for this thread."""
        self.ensure_one()

        try:
            query = "SELECT pg_try_advisory_lock(%s)"
            self.env.cr.execute(query, (self.id,))
            result = self.env.cr.fetchone()

            if not result or not result[0]:
                raise UserError(
                    _("Thread is currently generating a response. Please wait.")
                )

            _logger.info(f"Acquired advisory lock for thread {self.id}")

        except UserError:
            raise
        except OperationalError as e:
            _logger.error(f"Database error acquiring lock for thread {self.id}: {e}")
            raise UserError(_("Database error acquiring thread lock.")) from e
        except Exception as e:
            _logger.error(f"Unexpected error acquiring lock for thread {self.id}: {e}")
            raise UserError(_("Failed to acquire thread lock.")) from e

    def _release_thread_lock(self):
        """Release PostgreSQL advisory lock for this thread."""
        self.ensure_one()

        try:
            query = "SELECT pg_advisory_unlock(%s)"
            self.env.cr.execute(query, (self.id,))
            result = self.env.cr.fetchone()

            success = result and result[0]
            if success:
                _logger.info(f"Released advisory lock for thread {self.id}")
            else:
                _logger.warning(f"Advisory lock for thread {self.id} was not held")

            return success

        except Exception as e:
            _logger.error(f"Error releasing lock for thread {self.id}: {e}")
            return False

    @contextlib.contextmanager
    def _generation_lock(self):
        """Context manager for thread generation with automatic lock cleanup."""
        self.ensure_one()

        self._acquire_thread_lock()

        try:
            _logger.info(f"Starting locked generation for thread {self.id}")
            yield self

        finally:
            released = self._release_thread_lock()
            if released:
                _logger.info(f"Finished locked generation for thread {self.id}")
            else:
                _logger.warning(f"Lock release failed for thread {self.id}")

    # ============================================================================
    # ODOO HOOKS AND CLEANUP
    # ============================================================================

    @api.ondelete(at_uninstall=False)
    def _unlink_llm_thread(self):
        """Clean up ChromaDB collections when thread is deleted"""
        for thread in self:
            try:
                # Clean up ChromaDB collection for this thread
                rag_service = self.env['llm.rag.service']
                collection_name = f"thread_{thread.id}"
                
                client = rag_service._get_chromadb_client()
                if client:
                    try:
                        client.delete_collection(name=collection_name)
                        _logger.info(f"Deleted ChromaDB collection: {collection_name}")
                    except Exception as e:
                        _logger.warning(f"Failed to delete ChromaDB collection {collection_name}: {e}")
                        
            except Exception as e:
                _logger.error(f"Error cleaning up ChromaDB for thread {thread.id}: {e}")
        
        unlink_ids = [record.id for record in self]
        self.env["bus.bus"]._sendone(
            self.env.user.partner_id, "llm.thread/delete", {"ids": unlink_ids}
        )  

    # ============================================================================
    # HELPER METHODS FOR CONTEXT MANAGEMENT & UTILITIES
    # ============================================================================

    def _store_classification_context(self, classification_result, detected_tables):
        """Store classification context in a persistent JSON field."""
        self.ensure_one()
        context_data = {
            'classification_result': classification_result,
            'detected_tables': detected_tables,
            'timestamp': datetime.now().isoformat()
        }
        self.write({'pending_context': context_data})
        _logger.info(f"Stored pending context on thread {self.id}")

    def _get_classification_context(self):
        """Get stored classification context from the persistent JSON field."""
        self.ensure_one()
        context = self.pending_context
        if not context or not isinstance(context, dict):
            return None
            
        # Check if context is too old (expire after 30 minutes)
        stored_time_str = context.get('timestamp')
        if not stored_time_str:
            return None

        stored_time = datetime.fromisoformat(stored_time_str)
        if (datetime.now() - stored_time).total_seconds() > 1800:  # 30 minutes
            _logger.warning(f"Pending context for thread {self.id} has expired.")
            self._clear_classification_context() # Clear the expired context
            return None
            
        _logger.info(f"Retrieved valid pending context from thread {self.id}")
        return context

    def _clear_classification_context(self):
        """Clear the stored classification context from the persistent field."""
        self.ensure_one()
        self.write({'pending_context': None})
        _logger.info(f"Cleared pending context on thread {self.id}")

    def _resolve_table_name(self, table_name):
        """Resolve user input to actual table key"""
        table_name = table_name.lower().strip()
        
        # Direct match
        if table_name in self.SUPPORTED_TABLES:
            return table_name
            
        # Match by display name
        for key, config in self.SUPPORTED_TABLES.items():
            if config['name'].lower() == table_name:
                return key
            if table_name in config['name'].lower():
                return key
            if key.replace('_', ' ') == table_name.replace('_', ' '):
                return key
        
        return None

    def _get_thread_attachments(self):
        """Helper method to get all thread attachments"""
        self.env['ir.attachment'].invalidate_cache()
        self.env['mail.message'].invalidate_cache()
        
        thread_attachments = self.env['ir.attachment'].search([
            ('res_model', '=', 'llm.thread'),
            ('res_id', '=', self.id)
        ])
        
        messages = self.env['mail.message'].search([
            ('model', '=', 'llm.thread'),
            ('res_id', '=', self.id)
        ])
        
        message_attachments = self.env['ir.attachment'].search([
            ('res_model', '=', 'mail.message'),
            ('res_id', 'in', messages.ids)
        ]) if messages else self.env['ir.attachment']
        
        return thread_attachments | message_attachments

class OdooTableClassifier:
    """
    Enhanced classifier for determining appropriate Odoo table for data appending.
    This is the final, corrected version with a simplified and robust scoring system.
    """
    
    def __init__(self):
        # More comprehensive keyword mappings with weighted categories
        self.table_keywords = {
            'sale_order': {
                'primary': ['sales', 'sell', 'selling', 'customer', 'client', 'quotation', 'quote', 'revenue'],
                'secondary': ['order', 'so', 'customer order', 'sold', 'billing', 'invoice'],
                'context': ['price', 'amount', 'total', 'discount', 'tax', 'buyer', 'payment terms'],
                'columns': ['customer_name', 'product_name', 'quantity', 'unit_price', 'total_amount', 'order_date', 'partner_id']
            },
            'purchase_order': {
                'primary': ['purchase', 'buy', 'buying', 'vendor', 'supplier', 'procurement'],
                'secondary': ['po', 'purchase order', 'vendor order', 'bought', 'sourcing'],
                'context': ['cost', 'expense', 'delivery date', 'receipt', 'vendor_name'],
                'columns': ['vendor_name', 'product_name', 'quantity', 'unit_cost', 'total_cost', 'delivery_date', 'partner_id']
            },
            'stock_move': {
                'primary': ['stock movement', 'inventory movement', 'product movement', 'transfer'],
                'secondary': ['move', 'relocate', 'shift stock', 'internal transfer'],
                'context': ['quantity moved', 'source location', 'destination location'],
                'columns': ['product_name', 'quantity', 'from_location', 'to_location', 'move_date', 'product_id', 'location_id', 'location_dest_id']
            },
            'hr_expense': {
                'primary': ['expense', 'expenditure', 'spending', 'reimbursement', 'claim'],
                'secondary': ['business expense', 'travel expense', 'employee cost'],
                'context': ['employee name', 'expense date', 'receipt', 'category'],
                'columns': ['employee_name', 'expense_date', 'amount', 'category', 'description', 'employee_id']
            },
            # NOTE: We only need keywords for the tables we want to classify.
            # Other tables from SUPPORTED_TABLES will be handled but won't be auto-detected without keywords.
        }

    def analyze_user_intent(self, user_message, detected_tables):
        """
        Analyzes user intent using a simplified scoring system and returns normalized confidence.
        """
        user_message_lower = user_message.lower().strip()

        # Initialize scoring with a single 'raw_score' for simplicity and robustness
        classification_results = {}
        for table in self.table_keywords.keys():
            classification_results[table] = {'raw_score': 0}

        # Step 1: Score based on keywords in the user's message
        self._score_keywords(user_message_lower, classification_results)

        # Step 2: Score based on matching columns in the detected document tables
        self._analyze_table_structure(detected_tables, classification_results)

        # Sort results based on the final calculated raw_score
        sorted_results = sorted(
            classification_results.items(),
            key=lambda x: x[1]['raw_score'],
            reverse=True
        )

        # Normalize all scores against the total to get a true confidence percentage (0.0 to 1.0)
        total_raw_score = sum(item[1]['raw_score'] for item in sorted_results)
        if total_raw_score == 0:
            total_raw_score = 1  # Avoid division by zero if no matches are found

        for table, result in sorted_results:
            result['confidence'] = result['raw_score'] / total_raw_score

        primary_recommendation = sorted_results[0][0] if sorted_results and sorted_results[0][1]['raw_score'] > 0 else None
        primary_confidence = sorted_results[0][1]['confidence'] if primary_recommendation else 0.0

        # Return a clean, simple dictionary with the necessary results
        return {
            'primary_recommendation': primary_recommendation,
            'confidence': primary_confidence,
            'top_3_recommendations': [
                (table, result['confidence']) for table, result in sorted_results[:3]
            ],
        }

    def _score_keywords(self, user_message, classification_results):
        """Scores based on keyword matching and adds the result to the raw_score."""
        for table, keywords in self.table_keywords.items():
            score = 0
            for keyword in keywords.get('primary', []):
                if keyword in user_message: score += 5
            for keyword in keywords.get('secondary', []):
                if keyword in user_message: score += 3
            for keyword in keywords.get('context', []):
                if keyword in user_message: score += 1
            
            if table in classification_results:
                classification_results[table]['raw_score'] += score

    def _analyze_table_structure(self, detected_tables, classification_results):
        """Analyzes table column headers and adds the result to the raw_score."""
        for table_info in detected_tables:
            headers = [h.lower().strip().replace(' ', '_') for h in table_info.get('headers', [])]
            
            for odoo_table, keywords in self.table_keywords.items():
                if odoo_table not in classification_results: continue
                    
                column_matches = 0
                expected_columns = keywords.get('columns', [])
                if not expected_columns: continue

                for expected_col in expected_columns:
                    if any(expected_col in header for header in headers):
                        column_matches += 1
                
                match_percentage = column_matches / len(expected_columns)
                # A strong structural match is a very good indicator, so we weight it heavily
                classification_results[odoo_table]['raw_score'] += match_percentage * 10

    #
    # --- NOTE: The following methods from your original code have been intentionally REMOVED ---
    #
    # def _fuzzy_match(...)
    # def _analyze_context_intent(...)
    # def _apply_structure_bonuses(...)
    # def _generate_analysis_summary(...)
    #
    # These were part of the old, complex scoring system and were causing the KeyError.
    # Their logic has been simplified and integrated into the methods above.
    #


 
